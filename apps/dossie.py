import streamlit as st
import os
import json
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urljoin, urlparse
import markdown
from io import BytesIO
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from datetime import datetime
import logging
from loguru import logger
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import HumanMessage
from tenacity import retry, wait_fixed, retry_if_exception_type
from google.api_core.exceptions import DeadlineExceeded
from docx import Document
from docx.shared import Pt
import sys
from pathlib import Path

# Adiciona o diretório raiz ao path para importar módulos corretamente
sys.path.append(str(Path(__file__).parent.parent))

def app(config=None):
    # Verifica se a página já foi configurada pelo app principal
    if not config or not config.get("already_configured"):
        # Configuração da página Streamlit (só será executada se o app for executado sozinho)
        st.set_page_config(page_title='Gerador de Dossiê Comercial', layout="wide")
    
    # Configuração de logging
    log_dir = Path(__file__).parent.parent / "logs"
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / "dossie.log"
    
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(str(log_file)),
            logging.StreamHandler()
        ]
    )
    logger = logging.getLogger(__name__)

    # Carregamento de variáveis de ambiente
    try:
        from dotenv import find_dotenv, load_dotenv
        env_path = Path(__file__).parent.parent / ".env"
        load_dotenv(str(env_path))
        logger.info("Arquivo .env carregado com sucesso.")
    except Exception as e:
        logger.warning(f"Não foi possível carregar o arquivo .env: {e}")
        st.warning("Arquivo .env não encontrado. Algumas funcionalidades podem não estar disponíveis.")

    # Caminho para os assets
    assets_dir = Path(__file__).parent.parent / "assets"
    
    # Configuração do Gemini AI
    caminho_credenciais = os.getenv('GOOGLE_APPLICATION_CREDENTIALS', 
                                    str(Path(__file__).parent.parent / "decent-atlas-460512-g7-3b1d4ccb9c4e.json"))
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = caminho_credenciais

    modelo_ia = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.4)

    def custom_retry_decorator(wait_time=2):
        return retry(
            wait=wait_fixed(wait_time),
            retry=retry_if_exception_type(DeadlineExceeded),
            reraise=True
        )

    @custom_retry_decorator()
    def chat_with_retry(self, messages, **kwargs):
        start_time = time.time()
        while True:
            try:
                if time.time() - start_time > 180:  # 3 minutos
                    raise DeadlineExceeded("Timeout de 3 minutos excedido")
                return self.client.generate_content(messages, **kwargs)
            except DeadlineExceeded:
                raise  # Isso acionará o retry

    def normalizar_url(url):
        url = url.strip().lower()
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        parsed = urlparse(url)
        if not parsed.scheme:
            url = 'https://' + url
        if not parsed.netloc:
            return None
        
        return url

    def validar_url(url):
        try:
            response = requests.head(url, timeout=5)
            return response.status_code < 400
        except requests.RequestException:
            return False

    def obter_dominio(url):
        return urlparse(url).netloc

    def url_valida(url):
        analisada = urlparse(url)
        return bool(analisada.netloc) and bool(analisada.scheme)

    def obter_links_site(url):
        dominio = obter_dominio(url)
        urls = set()
        try:
            resposta = requests.get(url, timeout=10)
            resposta.raise_for_status()
            sopa = BeautifulSoup(resposta.content, "html.parser")
            for tag_a in sopa.find_all("a", href=True):
                href = tag_a['href']
                href_completo = urljoin(url, href)
                if url_valida(href_completo) and obter_dominio(href_completo) == dominio:
                    urls.add(href_completo)
            
            if not urls:
                urls = obter_links_com_selenium(url)
        except Exception as e:
            logger.error(f"Erro ao obter links de {url}: {e}")
            urls = obter_links_com_selenium(url)
        return urls

    def obter_links_com_selenium(url):
        urls = set()
        options = Options()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')

        try:
            service = Service(ChromeDriverManager().install())
            with webdriver.Chrome(service=service, options=options) as driver:
                driver.get(url)
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "a"))
                )
                links = driver.find_elements(By.TAG_NAME, "a")
                dominio = obter_dominio(url)
                for link in links:
                    href = link.get_attribute('href')
                    if href and url_valida(href) and obter_dominio(href) == dominio:
                        urls.add(href)
        except Exception as e:
            logger.error(f"Erro ao obter links com Selenium: {url} - {e}")
        return urls

    def raspar_com_selenium(url):
        options = Options()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-gpu')
        options.add_argument('--window-size=1920x1080')

        try:
            logger.debug(f"Iniciando raspagem com Selenium para: {url}")
            service = Service(ChromeDriverManager().install())
            with webdriver.Chrome(service=service, options=options) as driver:
                logger.debug("WebDriver Chrome iniciado com sucesso")
                driver.get(url)
                logger.debug(f"Página carregada: {driver.current_url}")
                
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                logger.debug("Elemento body encontrado na página")
                
                time.sleep(5)
                
                body = driver.find_element(By.TAG_NAME, "body")
                texto = body.text
                logger.debug(f"Texto extraído com sucesso. Tamanho: {len(texto)} caracteres")
                return texto
        except Exception as e:
            logger.error(f"Erro ao raspar com Selenium: {url} - {str(e)}", exc_info=True)
            return ""

    def extrair_contatos(sopa):
        contatos = {
            'telefones': set(),
            'emails': set(),
            'enderecos': set()
        }
        
        padrao_telefone = re.compile(r'($$?\d{2}$$?\s?(?:[2-8]|9[1-9])[0-9]{3}\-?[0-9]{4})')
        telefones = padrao_telefone.findall(str(sopa))
        contatos['telefones'].update(telefones)
        
        padrao_email = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        emails = padrao_email.findall(str(sopa))
        contatos['emails'].update(emails)
        
        enderecos = sopa.find_all('address')
        for endereco in enderecos:
            contatos['enderecos'].add(endereco.text.strip())
        
        return contatos

    def extrair_redes_sociais(sopa):
        redes_sociais = {}
        plataformas = ['facebook', 'twitter', 'instagram', 'linkedin', 'youtube']
        
        for plataforma in plataformas:
            links = sopa.find_all('a', href=re.compile(plataforma))
            if links:
                redes_sociais[plataforma] = links[0]['href']
        
        return redes_sociais

    def raspar_site(url, max_paginas=30):
        visitadas = set()
        para_visitar = {url}
        todo_texto = ""
        estrutura_site = {}
        palavras_chave = set()
        links_externos = set()
        tecnologias_usadas = set()
        contatos = {'telefones': set(), 'emails': set(), 'enderecos': set()}
        redes_sociais = {}
        produtos_servicos = set()
        equipe = set()
        parceiros = set()
        noticias = []

        while para_visitar and len(visitadas) < max_paginas:
            url_atual = para_visitar.pop()
            if url_atual not in visitadas:
                try:
                    logger.info(f"Raspando: {url_atual}")
                    resposta = requests.get(url_atual, timeout=10)
                    resposta.raise_for_status()
                    sopa = BeautifulSoup(resposta.content, 'html.parser')
                    
                    texto = sopa.get_text(separator=' ', strip=True)
                    if len(texto.split()) < 100:
                        texto = raspar_com_selenium(url_atual)
                    
                    texto = re.sub(r'\s+', ' ', texto)
                    todo_texto += texto + "\n\n"

                    estrutura_site[url_atual] = [h.text for h in sopa.find_all(['h1', 'h2', 'h3'])]

                    meta_keywords = sopa.find('meta', attrs={'name': 'keywords'})
                    if meta_keywords:
                        palavras_chave.update(meta_keywords['content'].split(','))

                    for link in sopa.find_all('a', href=True):
                        href = link['href']
                        if href.startswith('http') and obter_dominio(href) != obter_dominio(url):
                            links_externos.add(href)

                    if 'WordPress' in texto:
                        tecnologias_usadas.add('WordPress')
                    if 'Shopify' in texto:
                        tecnologias_usadas.add('Shopify')
                    if 'woocommerce' in str(sopa).lower():
                        tecnologias_usadas.add('WooCommerce')
                    if 'magento' in str(sopa).lower():
                        tecnologias_usadas.add('Magento')

                    novos_contatos = extrair_contatos(sopa)
                    for key in contatos:
                        contatos[key].update(novos_contatos[key])

                    redes_sociais.update(extrair_redes_sociais(sopa))

                    produtos = sopa.find_all('div', class_=re.compile('produto|servico'))
                    for produto in produtos:
                        produtos_servicos.add(produto.text.strip())

                    equipe_elementos = sopa.find_all('div', class_=re.compile('equipe|time|colaborador'))
                    for membro in equipe_elementos:
                        equipe.add(membro.text.strip())

                    parceiros_elementos = sopa.find_all('div', class_=re.compile('parceiro|cliente'))
                    for parceiro in parceiros_elementos:
                        parceiros.add(parceiro.text.strip())

                    noticias_elementos = sopa.find_all('article') or sopa.find_all('div', class_=re.compile('noticia|post'))
                    for noticia in noticias_elementos[:5]:
                        titulo = noticia.find('h2') or noticia.find('h3')
                        if titulo:
                            noticias.append(titulo.text.strip())

                    visitadas.add(url_atual)
                    novos_links = obter_links_site(url_atual)
                    para_visitar.update(novos_links - visitadas)
                    
                    logger.info(f"Total de texto coletado: {len(todo_texto)} caracteres")
                except Exception as e:
                    logger.error(f"Erro ao raspar {url_atual}: {e}")
                    texto = raspar_com_selenium(url_atual)
                    if texto:
                        todo_texto += texto + "\n\n"
                        visitadas.add(url_atual)

        logger.info(f"Total de páginas visitadas: {len(visitadas)}")
        logger.info(f"Tamanho final do texto coletado: {len(todo_texto)} caracteres")

        return {
            'texto_completo': todo_texto,
            'estrutura_site': estrutura_site,
            'palavras_chave': list(palavras_chave),
            'links_externos': list(links_externos),
            'tecnologias_usadas': list(tecnologias_usadas),
            'contatos': {k: list(v) for k, v in contatos.items()},
            'redes_sociais': redes_sociais,
            'produtos_servicos': list(produtos_servicos),
            'equipe': list(equipe),
            'parceiros': list(parceiros),
            'noticias': noticias
        }

    def gerar_dossie(dados_site, nome_empresa, url):
        prompt = f"""
        Você é um consultor sênior de negócios com MBA em processos comerciais, especializado em análise de empresas e estratégias de vendas. Com base nos dados fornecidos sobre a empresa {nome_empresa} (URL: {url}), crie um dossiê detalhado e estratégico. Este dossiê será usado pela equipe comercial para vender serviços de consultoria e pela equipe operacional para selecionar o perfil ideal de consultor para o projeto.

        Utilize as seguintes informações extraídas do site:

        Conteúdo do site: {dados_site['texto_completo'][:30000]}...

        Estrutura do site: {json.dumps(dados_site['estrutura_site'], indent=2)}

        Palavras-chave identificadas: {', '.join(dados_site['palavras_chave'])}

        Links externos relevantes: {', '.join(dados_site['links_externos'][:20])}

        Tecnologias utilizadas: {', '.join(dados_site['tecnologias_usadas'])}

        Contatos encontrados:
        - Telefones: {', '.join(dados_site['contatos']['telefones'])}
        - Emails: {', '.join(dados_site['contatos']['emails'])}
        - Endereços: {', '.join(dados_site['contatos']['enderecos'])}

        Redes Sociais: {json.dumps(dados_site['redes_sociais'], indent=2)}

        Produtos/Serviços: {', '.join(dados_site['produtos_servicos'])}

        Equipe: {', '.join(dados_site['equipe'])}

        Parceiros: {', '.join(dados_site['parceiros'])}

        Notícias recentes: {', '.join(dados_site['noticias'])}

        Crie um dossiê abrangente e estratégico com as seguintes seções, fornecendo análises aprofundadas, exemplos concretos, dados quantitativos (quando possível) e recomendações específicas para cada uma:

        1. Resumo Executivo
           - Apresente uma visão geral concisa e impactante da empresa, destacando seus principais pontos fortes, desafios e oportunidades.
           - Resuma as principais descobertas e recomendações do dossiê.
    Nome dos principais executivos da empresa:

        2. Perfil da Empresa
           - Detalhe a história e trajetória da empresa, utilizando todas as informações disponíveis e fazendo inferências lógicas quando necessário.
           - Identifique e analise a missão, visão e valores da empresa, mesmo que não estejam explicitamente declarados.
           - Estime o tamanho da empresa e sua presença no mercado, considerando fatores como variedade de produtos, alcance geográfico e presença online.

        3. Produtos e Serviços
           - Forneça uma descrição detalhada e categorizada da oferta de produtos e serviços.
           - Analise os principais diferenciadores de produto/serviço em relação à concorrência.
           - Identifique potenciais áreas para expansão ou melhoria do portfólio, baseando-se em tendências de mercado e necessidades dos clientes.

    Os produtos poderiam estar numa tabela?

        4. Análise de Mercado
           - Defina o público-alvo detalhadamente, incluindo demografia, comportamento de compra e necessidades específicas.
           - Analise o posicionamento da empresa no mercado, considerando fatores como preço, qualidade e diferenciação.
           - Identifique e compare os principais concorrentes, destacando pontos fortes e fracos em relação à {nome_empresa}.  Numa tabela.

        5. Estratégia de Marketing e Vendas
           - Detalhe os canais de marketing utilizados e sua eficácia aparente.
           - Analise as estratégias de aquisição de clientes identificáveis e sugira melhorias.
           - Proponha novas estratégias de marketing e vendas, incluindo táticas de marketing digital, parcerias estratégicas e programas de fidelização.

        6. Presença Digital
           - Realize uma análise aprofundada da estrutura e conteúdo do site, incluindo usabilidade, SEO e conversão.
           - Avalie a presença e estratégia de conteúdo nas mídias sociais.
           - Sugira oportunidades específicas de melhoria na presença online, incluindo redesign do site, estratégia de conteúdo e campanhas de marketing digital.

        7. Análise SWOT Detalhada  ( em tabela)
           - Forças: Identifique e analise detalhadamente as vantagens competitivas claras da empresa.
           - Fraquezas: Aponte as áreas que necessitam de melhoria, fornecendo exemplos concretos e impactos no negócio.
           - Oportunidades: Explore potenciais áreas de crescimento ou expansão, considerando tendências de mercado e forças da empresa.
           - Ameaças: Analise os desafios do mercado e da concorrência, sugerindo estratégias de mitigação.

        8. Cultura Organizacional e Recursos Humanos
           - Analise os valores e cultura corporativa perceptíveis através do conteúdo do site e comunicações.
           - Identifique políticas de RH ou benefícios mencionados, inferindo a abordagem da empresa em relação aos colaboradores.
           - Avalie o ambiente de trabalho e satisfação dos funcionários, se houver informações disponíveis.

        9. Tecnologia e Inovação
           - Avalie o nível de adoção tecnológica da empresa, considerando as tecnologias identificadas no site.
           - Identifique iniciativas de inovação ou P&D, se houver.
           - Sugira potenciais áreas para implementação tecnológica que poderiam melhorar a eficiência operacional e a experiência do cliente.

        10. Principais Dores do Cliente
            - Identifique e analise as dores mais significativas dos clientes com base no conteúdo do site e ofertas de produtos/serviços.
            - Analise padrões e tendências nas necessidades dos clientes, relacionando-os com a oferta atual da empresa.
            - Proponha oportunidades de melhoria e inovação baseadas nas dores identificadas.

        11. Oportunidades de Consultoria
            - Identifique áreas específicas onde serviços de consultoria poderiam agregar valor significativo, detalhando o impacto potencial.
            - Analise os desafios identificados que nossa empresa poderia ajudar a resolver, propondo soluções concretas.
            - Sugira projetos potenciais baseados nas necessidades percebidas e dores dos clientes, incluindo objetivos, escopo e resultados esperados.

        12. Perfil Ideal do Consultor
            - Detalhe as características e habilidades necessárias para atender esta empresa, considerando suas necessidades específicas.
            - Especifique a experiência relevante que seria mais valorizada para este cliente.
            - Descreva o estilo de trabalho que melhor se adequaria à cultura da empresa e aos desafios identificados.

        13. Estratégia de Abordagem Recomendada
            - Elabore pontos-chave detalhados a serem enfatizados na proposta de venda, alinhados com as necessidades e desafios da empresa.
            - Antecipe objeções potenciais e forneça estratégias específicas para superá-las.
            - Desenvolva uma estratégia de follow-up e construção de relacionamento de longo prazo com o cliente.

        14. Informações de Contato e Próximos Passos
            - Liste todos os detalhes de contato relevantes encontrados.
            - Proponha ações específicas e cronograma para iniciar o engajamento com a empresa.

        Forneça um dossiê extremamente detalhado, estratégico e acionável, usando dados concretos, exemplos específicos e análises aprofundadas sempre que possível. Use formatação Markdown para estruturar o conteúdo, incluindo listas, ênfases, citações e tabelas quando apropriado. O objetivo é fornecer insights valiosos e imediatamente aplicáveis para as equipes comercial e operacional.
        """

        mensagem = HumanMessage(content=prompt)
        resposta = modelo_ia.invoke([mensagem])
        dossie_conteudo = resposta.content if hasattr(resposta, 'content') else str(resposta)

        return dossie_conteudo

    def salvar_dossie_docx(conteudo, nome_arquivo):
        # Criar a pasta 'dossies' se ela não existir
        pasta_dossies = Path(__file__).parent.parent / 'dossies'
        pasta_dossies.mkdir(exist_ok=True)
        
        # Caminho completo do arquivo
        caminho_completo = pasta_dossies / nome_arquivo
        
        doc = Document()
        
        styles = doc.styles
        style_normal = styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(11)

        html = markdown.markdown(conteudo)
        soup = BeautifulSoup(html, 'html.parser')

        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
            if element.name in ['h1', 'h2', 'h3']:
                doc.add_heading(element.text, level=int(element.name[1]))
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet' if element.name == 'ul' else 'List Number')

        doc.save(str(caminho_completo))
        return str(caminho_completo)

    # Interface do usuário Streamlit
    col1, col2, col3 = st.columns([1,2,1])

    with col2:
        try:
            logo_path = assets_dir / "LOGO SUCESSO EM VENDAS HORIZONTAL AZUL.png"
            st.image(str(logo_path), width=300, use_container_width=True)
        except FileNotFoundError:
            st.write("Logo não encontrada. Por favor, verifique o caminho da imagem.")

    css = """
    <style>
        .centered-title {
            text-align: center;
            padding: 20px 0;
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)
    st.markdown("<h1 class='centered-title'>Gerador de Dossiê Comercial</h1>", unsafe_allow_html=True)
    st.write("")

    # Barra lateral
    with st.sidebar:
        st.header("Informações da Empresa")
        nome_empresa = st.text_input("Nome da Empresa:", "")
        url = st.text_input("URL do site da empresa:")
        
        if st.button("Gerar Dossiê"):
            if not nome_empresa or not url:
                st.error("Por favor, preencha todas as informações necessárias.")
            else:
                st.session_state.gerar_dossie = True

    # Container principal
    if 'gerar_dossie' in st.session_state and st.session_state.gerar_dossie:
        with st.spinner("Gerando dossiê..."):
            try:
                url_normalizada = normalizar_url(url)
                if not url_normalizada:
                    st.error("A URL fornecida é inválida. Por favor, verifique e tente novamente.")
                    st.session_state.gerar_dossie = False
                elif not validar_url(url_normalizada):
                    st.error("Não foi possível acessar a URL fornecida. Verifique se o site está online e tente novamente.")
                    st.session_state.gerar_dossie = False
                else:
                    st.info(f"Analisando o site: {url_normalizada}")
                    logger.info(f"Iniciando raspagem para: {url_normalizada}")
                    
                    # Raspar o site oficial
                    dados_site = raspar_site(url_normalizada)
                    if not dados_site['texto_completo']:
                        st.warning("Não foi possível extrair conteúdo do site oficial. O dossiê pode estar incompleto.")
                    
                    # Gerar o dossiê
                    dossie_conteudo = gerar_dossie(dados_site, nome_empresa, url_normalizada)
                    
                    # Exibir o conteúdo do dossiê na tela
                    st.markdown("## Dossiê Gerado")
                    st.markdown(dossie_conteudo)
                    
                    # Salvar o dossiê como DOCX
                    nome_arquivo = f"Dossie_{nome_empresa.replace(' ', '_')}.docx"
                    caminho_dossie = salvar_dossie_docx(dossie_conteudo, nome_arquivo)
                    
                    # Botão para download do DOCX
                    with open(caminho_dossie, "rb") as file:
                        docx_bytes = file.read()
                    
                    st.download_button(
                        label="Download Dossiê (DOCX)",
                        data=docx_bytes,
                        file_name=nome_arquivo,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success(f"Dossiê gerado com sucesso para {nome_empresa}!")
            except Exception as e:
                logger.error(f"Erro ao gerar o dossiê: {str(e)}", exc_info=True)
                st.error(f"Ocorreu um erro ao gerar o dossiê. Por favor, tente novamente mais tarde ou contate o suporte.")

        st.session_state.gerar_dossie = False

    # Botão para voltar à página inicial (no final do arquivo)
    if st.sidebar.button("← Voltar para a página inicial", key="btn_voltar_dossie"):
        st.session_state.current_app = 'home'
        st.query_params["app"] = "home"
        st.rerun()

if __name__ == "__main__":
    app()