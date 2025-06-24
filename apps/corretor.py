import streamlit as st
import os
import json
import time
from datetime import datetime
import logging
from loguru import logger
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import HumanMessage
from tenacity import retry, wait_fixed, retry_if_exception_type
from google.api_core.exceptions import DeadlineExceeded
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import io
import re
import difflib
import pandas as pd
import sys
from pathlib import Path

# Adiciona o diretório raiz ao path para importar módulos corretamente
sys.path.append(str(Path(__file__).parent.parent))

def app(config=None):
    # Verifica se a página já foi configurada pelo app principal
    if not config or not config.get("already_configured"):
        # Configuração da página Streamlit (só será executada se o app for executado sozinho)
        st.set_page_config(page_title='Corretor Ortográfico de Documentos', layout="wide")

    # Configuração de logging
    log_dir = Path(__file__).parent.parent / "logs"
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / "corretor.log"
    
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(str(log_file)),
            logging.StreamHandler()
        ]
    )
    logger = logging.getLogger(__name__)

    # Inicializar session_state
    if 'resultado_processamento' not in st.session_state:
        st.session_state.resultado_processamento = None
    if 'documento_corrigido' not in st.session_state:
        st.session_state.documento_corrigido = None
    if 'nome_arquivo' not in st.session_state:
        st.session_state.nome_arquivo = None

    # Carregamento de variáveis de ambiente
    try:
        from dotenv import find_dotenv, load_dotenv
        env_path = Path(__file__).parent.parent / ".env"
        load_dotenv(str(env_path))
        logger.info("Arquivo .env carregado com sucesso.")
    except Exception as e:
        logger.warning(f"Não foi possível carregar o arquivo .env: {e}")
        st.warning("Arquivo .env não encontrado. Algumas funcionalidades podem não estar disponíveis.")

    # Caminho para os assets
    assets_dir = Path(__file__).parent.parent / "assets"

    # Configuração do Gemini AI
    caminho_credenciais = os.getenv('GOOGLE_APPLICATION_CREDENTIALS', 
                                    str(Path(__file__).parent.parent / "decent-atlas-460512-g7-3b1d4ccb9c4e.json"))
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = caminho_credenciais

    modelo_ia = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.1)  # Temperatura mais baixa para ser mais conservador

    # Lista de termos protegidos que não devem ser alterados
    TERMOS_PROTEGIDOS = [
        "APONTE", "Aponte", "aponte", "A PONTE", "a ponte", "A ponte", "d'A PONTE", "d'APONTE", "d'Aponte", "d'aponte"
    ]

    def custom_retry_decorator(wait_time=2):
        return retry(
            wait=wait_fixed(wait_time),
            retry=retry_if_exception_type(DeadlineExceeded),
            reraise=True
        )

    @custom_retry_decorator()
    def chat_with_retry(messages, **kwargs):
        start_time = time.time()
        while True:
            try:
                if time.time() - start_time > 180:  # 3 minutos
                    raise DeadlineExceeded("Timeout de 3 minutos excedido")
                return modelo_ia.invoke(messages, **kwargs)
            except DeadlineExceeded:
                raise  # Isso acionará o retry

    def extrair_texto_docx(docx_bytes):
        """Extrai o texto de um arquivo DOCX mantendo a estrutura do documento."""
        doc = Document(io.BytesIO(docx_bytes))
        
        # Estrutura para armazenar o texto e formatação
        documento_estrutura = []
        
        for i, paragrafo in enumerate(doc.paragraphs):
            if paragrafo.text.strip():
                estilo = {
                    'indice': i,
                    'texto': paragrafo.text,
                    'estilo': paragrafo.style.name,
                    'alinhamento': paragrafo.alignment
                }
                documento_estrutura.append(estilo)
        
        return documento_estrutura

    def encontrar_alteracoes(texto_original, texto_corrigido):
        """Encontra as alterações específicas feitas no texto."""
        alteracoes = []
        
        # Usar SequenceMatcher para encontrar diferenças precisas
        s = difflib.SequenceMatcher(None, texto_original, texto_corrigido)
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag != 'equal':  # se houver diferença
                alteracoes.append({
                    'tipo': tag,
                    'original': texto_original[i1:i2],
                    'corrigido': texto_corrigido[j1:j2]
                })
        
        return alteracoes

    def preservar_termos_protegidos(texto_original, texto_corrigido):
        """Preserva os termos protegidos no texto corrigido."""
        # Criar um dicionário para mapear os termos protegidos encontrados no texto original
        termos_encontrados = {}
        
        for termo in TERMOS_PROTEGIDOS:
            # Encontrar todas as ocorrências do termo no texto original
            for match in re.finditer(re.escape(termo), texto_original):
                # Guardar a posição e o termo exato
                start, end = match.span()
                contexto = texto_original[max(0, start-10):min(len(texto_original), end+10)]
                termos_encontrados[contexto] = termo
        
        # Se não encontrou nenhum termo protegido, retorna o texto corrigido sem alterações
        if not termos_encontrados:
            return texto_corrigido
        
        # Para cada termo encontrado, verificar se foi alterado no texto corrigido
        texto_final = texto_corrigido
        for contexto, termo in termos_encontrados.items():
            # Tentar encontrar o contexto no texto corrigido
            contexto_pattern = re.escape(contexto).replace(re.escape(termo), '(.*?)')
            matches = re.finditer(contexto_pattern, texto_final)
            
            for match in matches:
                # Se encontrou o contexto, mas o termo está diferente
                if match.group(1) != termo:
                    # Substituir pelo termo original
                    antes = texto_final[:match.start(1)]
                    depois = texto_final[match.end(1):]
                    texto_final = antes + termo + depois
        
        # Verificar e corrigir casos específicos como "a ponte" -> "aponte"
        texto_final = re.sub(r'([Aa])\s+([Pp][Oo][Nn][Tt][Ee])', r'\1\2', texto_final)
        texto_final = re.sub(r'd\'([Aa])\s+([Pp][Oo][Nn][Tt][Ee])', r'd\'\1\2', texto_final)
        
        # Garantir que "A PONTE" permaneça em maiúsculas se estiver assim no original
        if "A PONTE" in texto_original and "a ponte" in texto_final.lower():
            texto_final = re.sub(r'[Aa]\s*[Pp][Oo][Nn][Tt][Ee]', "A PONTE", texto_final)
        
        return texto_final

    def corrigir_paragrafo_com_ia(paragrafo):
        """Corrige um único parágrafo com a IA."""
        texto_original = paragrafo['texto']
        
        # Se o texto for muito curto, não há necessidade de correção
        if len(texto_original) < 5:
            paragrafo['texto_corrigido'] = texto_original
            paragrafo['alteracoes'] = []
            return paragrafo
        
        # Identificar palavras em maiúsculas no texto original para preservar
        palavras_maiusculas = re.findall(r'\b[A-Z]{2,}\b', texto_original)
        
        prompt = f"""
        Você é um corretor ortográfico profissional. Sua tarefa é APENAS corrigir erros ortográficos, gramaticais e de pontuação no texto abaixo.
        
        REGRAS IMPORTANTES:
        1. NÃO altere o significado ou conteúdo do texto
        2. NÃO adicione ou remova informações
        3. NÃO altere nomes próprios, URLs, emails ou termos técnicos
        4. Mantenha EXATAMENTE a mesma estrutura do texto original
        5. Corrija APENAS erros de ortografia, gramática e pontuação
        6. Retorne APENAS o texto corrigido, sem comentários ou explicações
        7. Se não houver erros, retorne o texto exatamente como está
        8. PRESERVE maiúsculas e minúsculas exatamente como no texto original
        9. NUNCA altere a palavra "APONTE", "Aponte", "aponte", "A PONTE" ou qualquer variação - é um nome próprio
        10. Mantenha todas as palavras em MAIÚSCULAS exatamente como estão no original
        
        TERMOS QUE DEVEM SER PRESERVADOS EXATAMENTE COMO ESTÃO (não altere de forma alguma):
        - APONTE
        - Aponte
        - aponte
        - A PONTE
        - a ponte
        - d'A PONTE
        - d'APONTE
        - d'Aponte
        - d'aponte
        
        PALAVRAS EM MAIÚSCULAS QUE DEVEM PERMANECER EM MAIÚSCULAS:
        {', '.join(palavras_maiusculas)}
        
        TEXTO PARA CORREÇÃO:
        {texto_original}
        
        TEXTO CORRIGIDO:
        """
        
        mensagem = HumanMessage(content=prompt)
        
        try:
            resposta = chat_with_retry([mensagem])
            texto_corrigido = resposta.content if hasattr(resposta, 'content') else str(resposta)
            
            # Remover possíveis textos adicionais que não fazem parte da correção
            texto_corrigido = texto_corrigido.strip()
            if "TEXTO CORRIGIDO:" in texto_corrigido:
                texto_corrigido = texto_corrigido.split("TEXTO CORRIGIDO:", 1)[1].strip()
            
            # Verificar se o texto corrigido é muito diferente do original
            # Usando a similaridade de Levenshtein para garantir que não houve mudanças drásticas
            similaridade = calcular_similaridade(texto_original, texto_corrigido)
            
            # Se a similaridade for muito baixa, manter o texto original
            if similaridade < 0.7:  # Ajuste este valor conforme necessário
                logger.warning(f"Correção rejeitada por baixa similaridade ({similaridade:.2f}): {texto_original} -> {texto_corrigido}")
                paragrafo['texto_corrigido'] = texto_original
                paragrafo['alteracoes'] = []
            else:
                # Preservar termos protegidos
                texto_corrigido = preservar_termos_protegidos(texto_original, texto_corrigido)
                
                # Restaurar palavras em maiúsculas
                for palavra in palavras_maiusculas:
                    texto_corrigido = re.sub(r'\b' + re.escape(palavra.lower()) + r'\b', palavra, texto_corrigido, flags=re.IGNORECASE)
                
                # Encontrar as alterações específicas
                alteracoes = encontrar_alteracoes(texto_original, texto_corrigido)
                
                paragrafo['texto_corrigido'] = texto_corrigido
                paragrafo['alteracoes'] = alteracoes
            
            return paragrafo
        
        except Exception as e:
            logger.error(f"Erro ao corrigir texto com IA: {str(e)}")
            # Em caso de erro, retornar o texto original
            paragrafo['texto_corrigido'] = texto_original
            paragrafo['alteracoes'] = []
            return paragrafo

    def calcular_similaridade(texto1, texto2):
        """Calcula a similaridade entre dois textos usando a distância de Levenshtein."""
        # Normalizar textos para comparação (remover espaços extras, etc.)
        texto1_norm = re.sub(r'\s+', ' ', texto1).strip().lower()
        texto2_norm = re.sub(r'\s+', ' ', texto2).strip().lower()
        
        # Calcular similaridade usando SequenceMatcher
        s = difflib.SequenceMatcher(None, texto1_norm, texto2_norm)
        return s.ratio()

    def criar_documento_corrigido(documento_estrutura, doc_original):
        """Cria um novo documento DOCX com o texto corrigido, mantendo a formatação original."""
        doc_corrigido = Document()
        
        # Copiar estilos do documento original
        for style in doc_original.styles:
            if style.name not in doc_corrigido.styles:
                try:
                    doc_corrigido.styles.add_style(style.name, style.type)
                except:
                    pass  # Ignorar estilos que não podem ser copiados
        
        # Adicionar parágrafos corrigidos
        for paragrafo_info in documento_estrutura:
            texto_corrigido = paragrafo_info.get('texto_corrigido', paragrafo_info['texto'])
            
            # Criar novo parágrafo
            p = doc_corrigido.add_paragraph()
            
            # Aplicar estilo se disponível
            try:
                p.style = paragrafo_info['estilo']
            except:
                pass  # Usar estilo padrão se o estilo não existir
            
            # Aplicar alinhamento
            if 'alinhamento' in paragrafo_info and paragrafo_info['alinhamento'] is not None:
                p.alignment = paragrafo_info['alinhamento']
            
            # Adicionar texto
            p.add_run(texto_corrigido)
        
        return doc_corrigido

    def processar_documento(docx_bytes):
        """Processa o documento DOCX completo."""
        try:
            # Carregar documento original
            doc_original = Document(io.BytesIO(docx_bytes))
            
            # Extrair estrutura do documento
            documento_estrutura = extrair_texto_docx(docx_bytes)
            
            # Total de parágrafos
            total_paragrafos = len(documento_estrutura)
            documento_corrigido = []
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Processar cada parágrafo individualmente
            for i, paragrafo in enumerate(documento_estrutura):
                status_text.text(f"Processando parágrafo {i+1} de {total_paragrafos}...")
                paragrafo_corrigido = corrigir_paragrafo_com_ia(paragrafo)
                documento_corrigido.append(paragrafo_corrigido)
                progress_bar.progress((i + 1) / total_paragrafos)
            
            status_text.text("Montando documento final...")
            
            # Criar documento corrigido
            doc_final = criar_documento_corrigido(documento_corrigido, doc_original)
            
            # Salvar em memória
            output = io.BytesIO()
            doc_final.save(output)
            output.seek(0)
            
            status_text.text("Documento processado com sucesso!")
            return output, documento_corrigido
        
        except Exception as e:
            logger.error(f"Erro ao processar documento: {str(e)}", exc_info=True)
            st.error(f"Ocorreu um erro ao processar o documento: {str(e)}")
            return None, None

    # CSS personalizado
    css = """
    <style>
        .centered-title {
            text-align: center;
            padding: 20px 0;
        }
        .correction-highlight {
            background-color: #FFEB3B;
            padding: 2px 4px;
            border-radius: 3px;
        }
        .alteracao {
            margin-bottom: 5px;
            padding: 5px;
            border-radius: 4px;
            background-color: #f0f0f0;
        }
        .alteracao-replace {
            border-left: 4px solid #2196F3;
        }
        .alteracao-delete {
            border-left: 4px solid #F44336;
        }
        .alteracao-insert {
            border-left: 4px solid #4CAF50;
        }
        .sidebar-header {
            font-size: 1.2em;
            font-weight: bold;
            margin-bottom: 15px;
        }
        .sidebar-info {
            margin-bottom: 20px;
            font-size: 0.9em;
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

    # Interface do usuário Streamlit
    # Logo e título na área principal
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        try:
            logo_path = assets_dir / "LOGO SUCESSO EM VENDAS HORIZONTAL AZUL.png"
            st.image(str(logo_path), width=200)
        except FileNotFoundError:
            st.write("Logo não encontrada. Por favor, verifique o caminho da imagem.")

    st.markdown("<h1 class='centered-title'>Corretor Ortográfico de Documentos</h1>", unsafe_allow_html=True)
    
    # Barra lateral - Instruções e Upload
    st.sidebar.markdown("<div class='sidebar-header'>Instruções de Uso</div>", unsafe_allow_html=True)
    st.sidebar.markdown(
        "<div class='sidebar-info'>"
        "Faça upload de um documento DOCX para correção ortográfica automática. "
        "O documento será processado pelo Gemini AI e devolvido com as correções aplicadas, "
        "mantendo a formatação original."
        "</div>", 
        unsafe_allow_html=True
    )
    
    # Upload na barra lateral
    uploaded_file = st.sidebar.file_uploader("Escolha um arquivo DOCX", type=["docx"], key="corretor_file_uploader")
    
    # Botão para iniciar correção na barra lateral
    iniciar_correcao = False
    if uploaded_file is not None:
        st.sidebar.write(f"Arquivo: **{uploaded_file.name}**")
        
        if not st.session_state.resultado_processamento:
            iniciar_correcao = st.sidebar.button("Iniciar Correção", key="btn_iniciar_correcao", use_container_width=True)
    
    # Botão para voltar à página inicial
    st.sidebar.markdown("---")
    if st.sidebar.button("← Voltar para a página inicial", key="btn_voltar_corretor", use_container_width=True):
        st.session_state.current_app = 'home'
        st.query_params["app"] = "home"
        st.rerun()
    
    # Área principal - Resultados e processamento
    if iniciar_correcao:
        with st.spinner("Processando documento..."):
            docx_bytes = uploaded_file.getvalue()
            resultado, documento_corrigido = processar_documento(docx_bytes)
            
            # Salvar resultados na session_state
            if resultado:
                st.session_state.resultado_processamento = resultado
                st.session_state.documento_corrigido = documento_corrigido
                st.session_state.nome_arquivo = uploaded_file.name
                
                # Recarregar a página para mostrar os resultados
                st.rerun()

    # Exibir resultados se disponíveis na session_state
    if st.session_state.resultado_processamento:
        st.success("Documento corrigido com sucesso!")
        
        # Nome do arquivo de saída
        nome_original = st.session_state.nome_arquivo
        nome_base = os.path.splitext(nome_original)[0]
        nome_saida = f"{nome_base}_corrigido.docx"
        
        # Botão para download
        st.download_button(
            label="Download do Documento Corrigido",
            data=st.session_state.resultado_processamento,
            file_name=nome_saida,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="btn_download_corrigido"
        )
        
        # Exibir comparação de parágrafos
        st.markdown("## Correções Realizadas")
        
        # Contador de correções
        total_paragrafos_corrigidos = 0
        total_alteracoes = 0
        
        for i, paragrafo in enumerate(st.session_state.documento_corrigido):
            texto_original = paragrafo['texto']
            texto_corrigido = paragrafo.get('texto_corrigido', texto_original)
            alteracoes = paragrafo.get('alteracoes', [])
            
            if texto_original != texto_corrigido:
                total_paragrafos_corrigidos += 1
                total_alteracoes += len(alteracoes)
                
                st.markdown(f"### Parágrafo {i+1}")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Original:**")
                    st.text(texto_original)
                
                with col2:
                    st.markdown("**Corrigido:**")
                    st.text(texto_corrigido)
                
                # Mostrar as alterações específicas
                if alteracoes:
                    st.markdown("**Alterações:**")
                    for j, alt in enumerate(alteracoes):
                        tipo = alt['tipo']
                        original = alt['original']
                        corrigido = alt['corrigido']
                        
                        if tipo == 'replace':
                            st.markdown(f"<div class='alteracao alteracao-replace'>Substituído: '{original}' → '{corrigido}'</div>", unsafe_allow_html=True)
                        elif tipo == 'delete':
                            st.markdown(f"<div class='alteracao alteracao-delete'>Removido: '{original}'</div>", unsafe_allow_html=True)
                        elif tipo == 'insert':
                            st.markdown(f"<div class='alteracao alteracao-insert'>Adicionado: '{corrigido}'</div>", unsafe_allow_html=True)
                
                st.markdown("---")
        
        # Exibir estatísticas
        if total_paragrafos_corrigidos > 0:
            st.markdown(f"**Total de parágrafos corrigidos:** {total_paragrafos_corrigidos}")
            st.markdown(f"**Total de alterações realizadas:** {total_alteracoes}")
        else:
            st.info("Não foram encontradas correções a serem feitas no documento.")
        
        # Botão para limpar e reiniciar
        if st.button("Processar Novo Documento", key="btn_novo_documento"):
            # Limpar session_state
            st.session_state.resultado_processamento = None
            st.session_state.documento_corrigido = None
            st.session_state.nome_arquivo = None
            # Recarregar a página
            st.rerun()
    else:
        # Mensagem inicial quando não há resultados
        if not uploaded_file:
            st.info("👈 Faça o upload de um documento DOCX na barra lateral para começar.")
        elif not iniciar_correcao and not st.session_state.resultado_processamento:
            st.info("👈 Clique em 'Iniciar Correção' na barra lateral para processar o documento.")

if __name__ == "__main__":
    app()