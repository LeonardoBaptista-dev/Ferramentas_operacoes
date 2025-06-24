import io
from io import BytesIO
import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import HumanMessage
import os
import json
import docx
import PyPDF2
import pandas as pd
import logging
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import requests
from tenacity import retry, wait_fixed, retry_if_exception_type
from google.api_core.exceptions import DeadlineExceeded
from bs4 import BeautifulSoup
import re
from urllib.parse import urljoin, urlparse
import textwrap
import markdown
import tempfile
import base64
from PIL import Image
import markdown2
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import zipfile
import yagmail
import shutil
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from google.api_core.exceptions import ResourceExhausted, DeadlineExceeded
import sys
from pathlib import Path

# Adiciona o diretório raiz ao path para importar módulos corretamente
sys.path.append(str(Path(__file__).parent.parent))

def app(config=None):
    # Verifica se a página já foi configurada pelo app principal
    if not config or not config.get("already_configured"):
        # Configuração da página Streamlit (só será executada se o app for executado sozinho)
        st.set_page_config(page_title='Gerador de Método de Vendas Personalizado', layout="wide")
    
    # Configuração de logging
    log_dir = Path(__file__).parent.parent / "logs"
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / "metodo_vendas.log"
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(str(log_file)),
            logging.StreamHandler()
        ]
    )
    logger = logging.getLogger(__name__)

    # Carregamento de variáveis de ambiente
    try:
        from dotenv import find_dotenv, load_dotenv
        env_path = Path(__file__).parent.parent / ".env"
        load_dotenv(str(env_path))
        logger.info("Arquivo .env carregado com sucesso.")
    except Exception as e:
        logger.warning(f"Não foi possível carregar o arquivo .env: {e}")
        st.warning("Arquivo .env não encontrado. Algumas funcionalidades podem não estar disponíveis.")

    # Caminho para os assets e materiais
    assets_dir = Path(__file__).parent.parent / "assets"
    materials_dir = Path(__file__).parent.parent / "materiais"
    materials_dir.mkdir(exist_ok=True)  # Cria a pasta se não existir
    
    # Configuração do Gemini AI
    caminho_credenciais = os.getenv('GOOGLE_APPLICATION_CREDENTIALS', 
                                    str(Path(__file__).parent.parent / "decent-atlas-460512-g7-3b1d4ccb9c4e.json"))
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = caminho_credenciais

    modelo_ia = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.4)

    def custom_retry_decorator():
        return retry(
            wait=wait_exponential(multiplier=1, min=4, max=60),  # Backoff exponencial
            stop=stop_after_attempt(5),  # Máximo de 5 tentativas
            retry=retry_if_exception_type((ResourceExhausted, DeadlineExceeded)),
            reraise=True
        )

    @custom_retry_decorator()
    def chat_with_retry(self, messages, **kwargs):
        start_time = time.time()
        while True:
            try:
                if time.time() - start_time > 180:  # 3 minutos
                    raise DeadlineExceeded("Timeout de 3 minutos excedido")
                return self.client.generate_content(messages, **kwargs)
            except (ResourceExhausted, DeadlineExceeded) as e:
                # Log the error for debugging purposes
                logger.warning(f"Retry triggered: {str(e)}")
                raise  # This will trigger the retry mechanism

    # Funções para web scraping e geração de dossiê
    def obter_dominio(url):
        return urlparse(url).netloc

    def url_valida(url):
        analisada = urlparse(url)
        return bool(analisada.netloc) and bool(analisada.scheme)

    def obter_links_site(url):
        dominio = obter_dominio(url)
        urls = set()
        try:
            resposta = requests.get(url, timeout=10)
            resposta.raise_for_status()
            sopa = BeautifulSoup(resposta.content, "html.parser")
            for tag_a in sopa.find_all("a", href=True):
                href = tag_a['href']
                href_completo = urljoin(url, href)
                if url_valida(href_completo) and obter_dominio(href_completo) == dominio:
                    urls.add(href_completo)
            
            if not urls:  # Se não encontrar links, tente com Selenium
                urls = obter_links_com_selenium(url)
        except Exception as e:
            logger.error(f"Erro ao obter links de {url}: {e}")
            urls = obter_links_com_selenium(url)
        return urls

    def obter_links_com_selenium(url):
        urls = set()
        options = Options()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')

        try:
            service = Service(ChromeDriverManager().install())
            with webdriver.Chrome(service=service, options=options) as driver:
                driver.get(url)
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "a"))
                )
                links = driver.find_elements(By.TAG_NAME, "a")
                dominio = obter_dominio(url)
                for link in links:
                    href = link.get_attribute('href')
                    if href and url_valida(href) and obter_dominio(href) == dominio:
                        urls.add(href)
        except Exception as e:
            logger.error(f"Erro ao obter links com Selenium: {url} - {e}")
        return urls

    def raspar_site(url, max_paginas=10):
        visitadas = set()
        para_visitar = {url}
        todo_texto = ""

        while para_visitar and len(visitadas) < max_paginas:
            url_atual = para_visitar.pop()
            if url_atual not in visitadas:
                try:
                    # Tentativa com BeautifulSoup
                    resposta = requests.get(url_atual, timeout=10)
                    resposta.raise_for_status()
                    sopa = BeautifulSoup(resposta.content, 'html.parser')
                    
                    texto = sopa.get_text(separator=' ', strip=True)
                    if len(texto.split()) < 100:  # Se o texto extraído for muito curto, tente com Selenium
                        texto = raspar_com_selenium(url_atual)
                    
                    texto = re.sub(r'\s+', ' ', texto)
                    todo_texto += texto + "\n\n"

                    visitadas.add(url_atual)
                    
                    novos_links = obter_links_site(url_atual)
                    para_visitar.update(novos_links - visitadas)
                except Exception as e:
                    logger.error(f"Erro ao raspar {url_atual}: {e}")
                    # Se falhar com BeautifulSoup, tente com Selenium
                    texto = raspar_com_selenium(url_atual)
                    if texto:
                        todo_texto += texto + "\n\n"
                        visitadas.add(url_atual)

        return todo_texto

    def raspar_com_selenium(url):
        options = Options()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')

        try:
            service = Service(ChromeDriverManager().install())
            with webdriver.Chrome(service=service, options=options) as driver:
                driver.get(url)
                
                # Esperar pelo carregamento do conteúdo
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                
                # Dar um tempo extra para carregamento de conteúdo dinâmico
                time.sleep(5)
                
                # Extrair o conteúdo
                body = driver.find_element(By.TAG_NAME, "body")
                return body.text
        except Exception as e:
            logger.error(f"Erro ao raspar com Selenium: {url} - {e}")
            return ""

    def gerar_pdf_dossie(conteudo):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        flowables = []

        # Converter Markdown para HTML
        html = markdown.markdown(conteudo)
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
            if element.name in ['h1', 'h2', 'h3']:
                style = styles['Heading' + element.name[1]]
                flowables.append(Paragraph(element.text, style))
            elif element.name == 'p':
                flowables.append(Paragraph(element.text, styles['Normal']))
            elif element.name in ['ul', 'ol']:
                for index, li in enumerate(element.find_all('li'), start=1):
                    bullet = '•' if element.name == 'ul' else f"{index}."
                    flowables.append(Paragraph(f"{bullet} {li.text}", styles['Normal']))
            flowables.append(Spacer(1, 12))

        doc.build(flowables)
        buffer.seek(0)
        return buffer

    def gerar_dossie(conteudo_site):
        prompt = f"""
        Com base no conteúdo do site fornecido abaixo, crie um dossiê detalhado da empresa. 
        O dossiê deve incluir as seguintes seções:

        1. Visão geral da empresa
        2. Produtos ou serviços principais
        3. Público-alvo
        4. Proposta de valor única
        5. Presença online e mídias sociais
        6. Análise SWOT (Forças, Fraquezas, Oportunidades, Ameaças)
        7. Concorrentes principais (se mencionados)
        8. Cultura e valores da empresa
        9. Informações de contato e localização

        Conteúdo do site:
        {conteudo_site}

        Por favor, forneça um dossiê estruturado e informativo baseado nas informações disponíveis.
        """
        
        mensagem = HumanMessage(content=prompt)
        resposta = modelo_ia.invoke([mensagem])
        dossie_conteudo = resposta.content if hasattr(resposta, 'content') else str(resposta)

        pdf_buffer = gerar_pdf_dossie(dossie_conteudo)
        return pdf_buffer

    def gerar_e_salvar_dossie(url, cliente_nome):
        conteudo_site = raspar_site(url)
        if conteudo_site:
            pdf_buffer = gerar_dossie(conteudo_site)
            
            nome_arquivo = f"Dossie_{cliente_nome.replace(' ', '_')}.pdf"
            caminho_dossie = os.path.join(tempfile.gettempdir(), nome_arquivo)
            
            with open(caminho_dossie, 'wb') as f:
                f.write(pdf_buffer.getvalue())
            
            return caminho_dossie
        return None

    # Funções para carregar diferentes tipos de arquivos
    def carregar_arquivo(arquivo, max_size=1000000):  # 1MB por padrão
        if isinstance(arquivo, str):  # Se for um caminho de arquivo
            if os.path.getsize(arquivo) > max_size:
                logger.warning(f"Arquivo {arquivo} é muito grande. Carregando apenas os primeiros {max_size} bytes.")
                with open(arquivo, 'rb') as f:
                    return f.read(max_size).decode('utf-8', errors='ignore')
            
            _, extensao_arquivo = os.path.splitext(arquivo)
            with open(arquivo, 'rb') as f:
                conteudo = f.read()
        else:  # Se for um objeto FileUploader do Streamlit
            conteudo = arquivo.getvalue()
            extensao_arquivo = os.path.splitext(arquivo.name)[1]

        try:
            if extensao_arquivo == '.csv':
                # Tenta diferentes codificações
                encodings = ['utf-8', 'iso-8859-1', 'windows-1252']
                for encoding in encodings:
                    try:
                        df = pd.read_csv(io.BytesIO(conteudo), encoding=encoding, on_bad_lines='skip')
                        return df.to_string()
                    except UnicodeDecodeError:
                        continue
                raise ValueError(f"Não foi possível decodificar o arquivo CSV com as codificações: {encodings}")
            elif extensao_arquivo == '.xlsx':
                # Salvar temporariamente o arquivo
                with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
                    tmp.write(conteudo)
                    tmp_path = tmp.name
                
                # Ler o arquivo Excel
                df = pd.read_excel(tmp_path, engine='openpyxl')
                
                # Remover o arquivo temporário
                os.unlink(tmp_path)
                
                return df.to_string()
            elif extensao_arquivo == '.json':
                return json.loads(conteudo)
            elif extensao_arquivo == '.docx':
                doc = Document(io.BytesIO(conteudo))
                return "\n".join([p.text for p in doc.paragraphs])
            elif extensao_arquivo == '.pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(conteudo))
                return "\n".join([page.extract_text() for page in pdf_reader.pages])
            elif extensao_arquivo in ['.md', '.txt']:
                # Tenta diferentes codificações
                encodings = ['utf-8', 'iso-8859-1', 'windows-1252']
                for encoding in encodings:
                    try:
                        return conteudo.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                raise ValueError(f"Não foi possível decodificar o arquivo de texto com as codificações: {encodings}")
            else:
                return conteudo.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Erro ao carregar o arquivo {arquivo.name if hasattr(arquivo, 'name') else arquivo}: {e}")
            return ""

    # Função para carregar todos os arquivos de uma pasta
    def carregar_conteudo_pasta(caminho_pasta):
        conteudos = []
        if not os.path.exists(caminho_pasta):
            logger.warning(f"Pasta {caminho_pasta} não encontrada")
            return ""
            
        for nome_arquivo in os.listdir(caminho_pasta):
            caminho_arquivo = os.path.join(caminho_pasta, nome_arquivo)
            if os.path.isfile(caminho_arquivo):
                conteudo = carregar_arquivo(caminho_arquivo)
                if conteudo:
                    conteudos.append(conteudo)
        return "\n\n".join(conteudos)

    # Função para gerar o método de vendas completo
    def gerar_metodo_vendas_completo(conteudo_arquivos, materiais, modelo_metodo_vendas, opcoes_personalizacao):
        partes = [
            "INTRODUÇÃO",
            "O que é o Método de Vendas A PONTE",
            "ABORDE POSITIVAMENTE",
            "PESQUISE O CLIENTE",
            "OFERECA SOLUCÕES",
            "NEGOCIE E NEUTRALIZE AS OBJEÇÕES",
            "TOME A INICIATIVA E FECHE A VENDA",
            "ESTENDA O RELACIONAMENTO",
            "CONCLUSÃO",
            "CADERNO DE OBJEÇÕES",
            "CADERNO DE BENEFÍCIOS",
            "RAIO-X",
            "DE VOLTA A LOJA",
            "MANUAL DE BOLSO"
        ]
        
        metodo_completo = ""
        for parte in partes:
            st.write(f"Gerando seção: {parte}")
            parte_gerada = gerar_parte_metodo(parte, conteudo_arquivos, materiais, modelo_metodo_vendas, opcoes_personalizacao)
            metodo_completo += f"# {parte}\n\n{parte_gerada}\n\n"
        
        return metodo_completo

    def gerar_diretrizes_linguagem(tipos_linguagem):
        diretrizes = "Clareza e Adaptação: "
        if "Formal" in tipos_linguagem:
            diretrizes += "Use uma linguagem formal e profissional. "
        if "Informal" in tipos_linguagem:
            diretrizes += "Adote um tom mais casual e amigável. "
        if "Técnico" in tipos_linguagem:
            diretrizes += "Incorpore termos técnicos relevantes, mas sempre com explicações claras. "
        if "Consultivo" in tipos_linguagem:
            diretrizes += "Adote um tom de aconselhamento profissional. "
        if "Popular" in tipos_linguagem:
            diretrizes += "Use uma linguagem simples e acessível. "
        diretrizes += "Adapte o conteúdo ao público-alvo, mantendo a clareza e relevância para todos os níveis."
        return diretrizes

    def gerar_instrucoes_canais(canais_atendimento):
        instrucoes = "Foque nas técnicas de venda específicas para os seguintes canais: "
        for canal in canais_atendimento:
            if canal == "WhatsApp":
                instrucoes += "WhatsApp (enfatize comunicação escrita clara, uso de recursos visuais e áudio, e resposta rápida). "
            elif canal == "Telefone":
                instrucoes += "Telefone (destaque a importância do tom de voz, escuta ativa e scripts eficientes). "
            elif canal == "Presencial":
                instrucoes += "Presencial (aborde linguagem corporal, demonstração de produtos e criação de rapport). "
            elif canal == "PaP":
                instrucoes += "Porta a Porta (foque em preparação prévia, adaptação ao ambiente do cliente e técnicas de apresentação móveis). "
        return instrucoes

    def gerar_instrucoes_tamanho(tamanho_material):
        if tamanho_material == "Curto":
            return "Mantenha o conteúdo conciso e direto ao ponto, focando nos elementos mais essenciais. Priorize exemplos práticos e dicas que possam ser implementadas."
        elif tamanho_material == "Médio":
            return "Forneça um equilíbrio entre detalhes e concisão, cobrindo os pontos principais com explicações moderadas. Inclua exemplos práticos, tabelas e algumas análises mais aprofundadas, mantendo o foco na aplicabilidade."
        elif tamanho_material == "Longo":
            return "Forneça explicações detalhadas, exemplos extensos e informações contextuais aprofundadas. Explore cada conceito em profundidade, oferecendo múltiplos exemplos, tabelas, casos de estudo e análises detalhadas. Inclua  seções de 'aprofundamento' para tópicos mais complexos."
        else:
            return "Adapte o conteúdo de forma equilibrada, fornecendo informações suficientes para uma compreensão completa, mas mantendo o foco na aplicação prática."

    def gerar_instrucoes_segmento(segmento):
        instrucoes = f"Adapte o conteúdo especificamente para o segmento de {segmento}. "
        if segmento == "Concessionárias de veículos":
            instrucoes += "Foque em técnicas de venda para automóveis, ciclo de compra de veículos e peculiaridades do mercado automotivo. "
        elif segmento == "Imóveis":
            instrucoes += "Aborde estratégias para vendas de longo prazo, financiamento imobiliário e aspectos legais da venda de imóveis. "
        elif segmento == "Eletromóveis":
            instrucoes += "Enfatize conhecimento técnico de produtos, venda de garantias estendidas e ciclos de renovação de eletrodomésticos. "
        elif segmento == "Serviços Financeiros":
            instrucoes += "Foque em compliance, explicação de produtos financeiros complexos e construção de confiança com o cliente. "
        elif segmento == "Cama, Mesa e Banho":
            instrucoes += "Destaque técnicas de venda cruzada, conhecimento de materiais e tendências de decoração. "
        elif segmento == "Tintas":
            instrucoes += "Aborde conhecimento técnico de produtos, combinação de cores e soluções para diferentes superfícies e ambientes. "
        elif segmento == "Farma":
            instrucoes += "Foque em regulamentações do setor farmacêutico, ética na venda de medicamentos e conhecimento de produtos de saúde. "
        elif segmento == "Mat. Construção":
            instrucoes += "Enfatize conhecimento técnico de materiais, soluções para projetos de construção e reforma, e vendas para profissionais e consumidores finais. "
        else:
            instrucoes += "Adapte as técnicas de venda para as especificidades deste segmento, considerando o ciclo de compra, perfil do consumidor e particularidades do produto/serviço."
        return instrucoes

    def gerar_parte_metodo(parte, conteudo_arquivos, materiais, modelo_metodo_vendas, opcoes_personalizacao):
        # Definir prompts específicos para cada parte
        prompts = {
        "INTRODUÇÃO": """
        Crie uma introdução abrangente para o método de vendas.
        Comece com uma mensagem de entrada semelhante a essa: 
        
        "Olá, Profissional de Vendas!
        Seja bem-vindo ao metodo de vendas da (nome empresa) (segmento da empresa)
        Este material é uma ferramenta prática e dinâmica para ajudar e orientar o profissional de vendas da (nome empresa) (segmento da empresa) a obter
        melhores resultados.
        Temos certeza que, se você souber usufruir todos os benefícios existentes
        neste material e de todas as atividades que serão desenvolvidas daqui em
        diante, você obterá muito sucesso em suas vendas.
        Para que este conteúdo cumpra com a sua eficácia, é necessário que
        ele seja parte integrante do material de trabalho e, também, aplicado
        diariamente em todos os atendimentos pelo time de vendas da (nome empresa) (segmento da empresa).
        
        "Vendas é ciência e deve ser tratada como tal."

        O que é vender?
        Nessa introdução não cite as etapas do metodo APONTE e nem exemplos praticos.
        """,

        "O QUE É O MÉTODO DE VENDAS A PONTE": """
        
        1. Origem e desenvolvimento do método
        2. Porque que os consultores precisam de um Método de vendas?
        3. O que o consultor ganha com o Método de Vendas
        4. Visão geral de cada etapa do APONTE
        5. Como o método se diferencia de outras abordagens de vendas
        6. Benefícios esperados da implementação do método
        7. Como usar o Método de Vendas?

        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        
        """,
        
        "ABORDE POSITIVAMENTE": """
        Detalhe a Etapa 'Aborde Positivamente' do Método A PONTE porém lembre-se de não ser redundante.

        Ensinar os vendedores a estabelecer uma conexão inicial positiva e acolhedora com os clientes, criando uma base sólida para o relacionamento e o processo de vendas subsequente.

        Diretrizes:

        Primeira Impressão é Fundamental:

        Enfatize a importância de causar uma boa primeira impressão. O vendedor deve se apresentar de maneira positiva e acolhedora, com um sorriso e expressões faciais abertas, independentemente do canal de comunicação.
        Demonstre como o uso de uma postura positiva, tanto física quanto vocal, pode impactar diretamente a percepção do cliente.
        Uso do Nome e Personalização:

        Destaque a importância de usar o nome do cliente durante a interação para personalizar o atendimento e criar uma conexão mais profunda.
        Pratique técnicas para lembrar nomes e detalhes importantes dos clientes, como a repetição e o uso de notas.
        Empatia e Sintonia com o Cliente:

        Instrua os vendedores a lerem e se adaptarem ao humor e às emoções do cliente. Isso inclui ajustar o nível de entusiasmo e formalidade com base no feedback não verbal e verbal do cliente.
        Ensine como empregar empatia ativamente, validando as emoções e as perspectivas do cliente.
        Construindo Confiança Através da Linguagem Corporal:

        Explore como a linguagem corporal pode ser usada para transmitir confiança e receptividade. Inclua práticas sobre postura, contato visual, e gestos que indicam atenção e respeito.
        Realce a importância de refletir sobre a própria apresentação e comportamento após cada interação, buscando sempre melhorias.
        Feedback e Melhoria Contínua:

        Estabeleça um processo para receber feedback regularmente, seja através de colegas, supervisores ou até mesmo do cliente, quando apropriado.
        Encoraje os vendedores a ajustarem suas abordagens com base no feedback e nas experiências de vendas passadas, promovendo um ciclo contínuo de aprendizado e aprimoramento.
         Detalhe a Etapa 'Aborde Positivamente' do Método A PONTE seguindo ESTRITAMENTE a estrutura abaixo. Cada seção deve ser claramente identificada e desenvolvida:

        1. Objetivo:
        [Desenvolva este tópico]

        2. O que fazer:
        [Desenvolva este tópico]

        3. Perfis de clientes:
        [Desenvolva este tópico]

        4. Abordagem:
        [Desenvolva este tópico]

        5. Como fazer uma Abordagem Positiva:
        [Desenvolva este tópico]

        6. O que o vendedor (a) ganha em abordar positivamente?
        [Desenvolva este tópico]

        7. O que o cliente ganha com a abordagem positiva?
        [Desenvolva este tópico]

        8. Como abordar positivamente influencia na prospecção de clientes:
        [Desenvolva este tópico]

        9. Por que o vendedor precisa abordar positivamente?
        [Desenvolva este tópico]

        IMPORTANTE: Mantenha-se fiel a esta estrutura, desenvolvendo cada tópico separadamente e na ordem apresentada. Não adicione seções extras nem omita nenhuma das seções solicitadas.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.

        """,

        "PESQUISE O CLIENTE": """

        Detalhe a Etapa 'Pesquise o Cliente' do Método A PONTE porém lembre-se de não ser redundante.


        Capacitar os vendedores a conduzir uma pesquisa eficaz sobre os clientes, utilizando perguntas estratégicas para entender profundamente suas necessidades, preferências, e potencial financeiro, visando construir relacionamentos fortes e maximizar as oportunidades de vendas.

        Diretrizes:

        Entendimento Profundo do Cliente:

        Enfatize a importância de conhecer o cliente além das necessidades superficiais. Instrua os vendedores a identificar gostos, estilos, comportamentos, valores, e potenciais financeiros dos clientes.
        Oriente sobre como adaptar as perguntas e a abordagem ao contexto específico do cliente, usando informações de perfis para personalizar a interação.
        Técnicas de Perguntas:

        Ensine o uso de perguntas abertas para iniciar diálogos que revelam informações profundas sobre os clientes. Exemplos de perguntas abertas incluem: "O que o traz à nossa loja hoje?" ou "Como você imagina o produto ideal para atender às suas necessidades?"
        Explique quando e como utilizar perguntas fechadas para clarificar pontos específicos ou confirmar entendimentos.
        Escuta Ativa e Observação:

        Ressalte a necessidade de escutar mais do que falar. Instrua os vendedores a observar sinais não verbais como expressões faciais, gestos e o tom de voz para ajustar a abordagem em tempo real.
        Encoraje a utilização de técnicas de escuta ativa, como acenos e frases que demonstrem compreensão, para fazer com que o cliente se sinta ouvido e compreendido.
        
        Detalhe a Etapa 'Pesquise o Cliente' do Método A PONTE, seguindo ESTRITAMENTE a estrutura abaixo. Cada seção deve ser claramente identificada e desenvolvida de forma extensa:

        1. Objetivo:
        [Desenvolva este tópico detalhadamente]

        2. O que o consultor ganha quando pesquisa o cliente?
        [Desenvolva este tópico detalhadamente]

        3. O que o cliente ganha com uma pesquisa assertiva?
        [Desenvolva este tópico detalhadamente]

        4. Por que o consultor precisa pesquisar?
        [Desenvolva este tópico detalhadamente]

        5. Quantas perguntas se deve fazer para o cliente?
        [Desenvolva este tópico detalhadamente]

        6. Potencial de compra / capacidade financeira:
        [Desenvolva este tópico detalhadamente]

        7. Motivação / emocional:
        [Desenvolva este tópico detalhadamente]

        8. Senso de urgência:
        [Desenvolva este tópico detalhadamente]

        9. Passos da pesquisa:
        a) Peça licença e pergunte o que quiser
        b) O que perguntar?
        c) Perguntas Fechadas são ruins?
        d) Necessidades presentes e futuras
        [Desenvolva cada um destes subtópicos detalhadamente]

        10. Conhecendo nosso cliente:
        Faça perguntas abertas e focadas no conhecimento das motivações de compra do cliente.
        [Desenvolva este tópico detalhadamente]

        11. Integração de Tecnologia:
        Explique como utilizar o CRM para registrar e acessar informações do cliente durante esta etapa.
        [Desenvolva este tópico detalhadamente]

        IMPORTANTE: 
        - Mantenha-se fiel a esta estrutura, desenvolvendo cada tópico separadamente e na ordem apresentada. 
        - Não adicione seções extras nem omita nenhuma das seções solicitadas.
        - Certifique-se de que cada seção seja claramente identificada com seu número e título correspondente.
        - Lembre-se de adaptar o conteúdo para o contexto da empresa.
        - Evite redundâncias entre as seções, mas mantenha a coesão do texto como um todo.
        
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        Ao elaborar o conteúdo, considere as diretrizes fornecidas sobre entendimento profundo do cliente, técnicas de perguntas, escuta ativa e observação, análise e utilização das informações, e integração com etapas subsequentes do método A PONTE.
        """,

        "OFERECA SOLUCÕES": """

        Detalhe a Etapa 'Oferecer uma Demonstração Envolvente' do Método A PONTE porém lembre-se de não ser redundante.

        Capacitar os vendedores a realizar demonstrações que não apenas informem, mas envolvam e emocionem os clientes, destacando os benefícios dos produtos ou serviços de forma a minimizar objeções e acelerar o fechamento das vendas.

        Diretrizes:

        Preparação Detalhada:

        Ensine os vendedores a se prepararem minuciosamente para as demonstrações, garantindo um profundo conhecimento sobre os produtos ou serviços que estão vendendo, incluindo características técnicas e benefícios.
        Instrua-os sobre a importância de adaptar a demonstração para ressoar com os interesses e necessidades específicos do cliente, baseados nas informações coletadas na fase de pesquisa.
        Demonstração dos Benefícios:

        Oriente os vendedores a focar nos benefícios em vez de apenas nas características. Cada característica deve ser vinculada a um benefício claro que resolva um problema do cliente ou melhore sua situação de alguma forma.
        Encoraje-os a formular perguntas como "O que o cliente ganha com isso?" para cada característica, para garantir que os benefícios sejam comunicados de forma eficaz e emocionalmente engajante.
        Engajamento e Interatividade:

        Aconselhe os vendedores a envolverem os clientes na demonstração, permitindo-lhes experimentar o produto ou serviço quando possível.
        Mostre como usar perguntas durante a demonstração para manter o cliente engajado e abrir caminho para uma comunicação bidirecional, permitindo que o vendedor ajuste sua abordagem conforme a reação do cliente.
        Oferta de Produtos e Serviços Complementares:

        Prepare os vendedores para aumentar o ticket médio, apresentando produtos ou serviços complementares e suplementares que agreguem valor à compra principal do cliente.
        Ensine técnicas para introduzir essas ofertas de forma natural e relevante, destacando como esses adicionais podem enriquecer ainda mais a experiência do cliente ou resolver outros problemas potenciais que eles possam ter.
        Fechamento da Demonstração:

        Instrua sobre como concluir a demonstração de maneira forte, resumindo os principais benefícios discutidos e reafirmando como o produto ou serviço pode atender ou superar as expectativas do cliente.
        Encoraje os vendedores a fazerem perguntas finais para verificar a compreensão e satisfação do cliente antes de moverem-se para o fechamento da venda.
        Dê exemplos práticos.
  
        Detalhe a Etapa 'Ofereça Soluções' do Método A PONTE, seguindo ESTRITAMENTE a estrutura abaixo. Cada seção deve ser claramente identificada e desenvolvida de forma extensa:

        1. Objetivo:
        [Desenvolva este tópico detalhadamente]

        2. O que fazer:
        [Desenvolva este tópico detalhadamente]

        3. O que é oferecer soluções:
        [Desenvolva este tópico detalhadamente]

        4. Características e benefícios:
        [Desenvolva este tópico detalhadamente]

            
        5. Storytelling / demonstração envolvente:
        [Desenvolva este tópico detalhadamente, incluindo exemplos de como aplicar storytelling na venda]

        6. O que o cliente ganha com isso?:
        [Desenvolva este tópico detalhadamente]

        7. Tabela Característica x Benefícios:
        [Crie uma tabela detalhada com características dos produtos da empresa e seus respectivos benefícios]

        8. Como oferecer soluções?:
        [Desenvolva este tópico detalhadamente, fornecendo passos práticos]

        


        IMPORTANTE: 
        - Mantenha-se fiel a esta estrutura, desenvolvendo cada tópico separadamente e na ordem apresentada. 
        - Não adicione seções extras nem omita nenhuma das seções solicitadas.
        - Certifique-se de que cada seção seja claramente identificada com seu número e título correspondente.
        - Evite redundâncias entre as seções, mas mantenha a coesão do texto como um todo.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        """,

        "NEGOCIE E NEUTRALIZE AS OBJEÇÕES" : """

        Detalhe a Etapa 'Negocie e Neutralize Objeções' do Método A PONTE porém lembre-se de não ser redundante.

        Capacitar os vendedores a gerenciar negociações e neutralizar objeções de maneira eficaz, utilizando técnicas apropriadas para garantir que tanto o cliente quanto o vendedor saiam da interação satisfeitos e com suas necessidades atendidas.

        Diretrizes:

        Preparação e Comportamento:

        Enfatize a importância de abordar a negociação com um comportamento calmo e receptivo. A negociação não deve ser vista como um confronto, mas como uma oportunidade de entender e atender melhor às necessidades do cliente.
        Treine os vendedores para ouvirem as objeções do cliente com paciência, sem interrupções, mostrando que estão genuinamente interessados em encontrar uma solução que beneficie ambas as partes.
        Técnicas de Resposta a Objeções:

        Introduza técnicas práticas para responder a objeções, como a estratégia de converter objeções em perguntas que levem o cliente a expressar suas verdadeiras preocupações ou expectativas.
        Instrua sobre o uso de empatia e validação nas respostas, como começar com "Eu entendo..." ou "Compreendo sua preocupação...", antes de oferecer uma solução ou alternativa.
        Demonstração de Valor e Benefícios:

        Mostre aos vendedores como enfatizar os benefícios dos produtos ou serviços durante a negociação, relacionando-os diretamente às necessidades e desejos expressos pelo cliente.
        Prepare-os para usar informações técnicas e benefícios de maneira estratégica, para reforçar o valor do que está sendo oferecido e diminuir a resistência do cliente.
        Prática com Exemplos e Role-playing:

        Utilize exemplos práticos e simulações de role-playing para treinar os vendedores sobre como aplicar as técnicas de negociação em situações reais. Isso inclui lidar com diferentes tipos de objeções, desde preocupações com preço até dúvidas sobre a adequação do produto.
        Encoraje a análise de casos de estudo ou vídeos de treinamento, como o mencionado exemplo do seriado Dr. House, para ilustrar como técnicas eficazes de comunicação e negociação podem ser aplicadas.
        Fechamento da Negociação:

        Ensine técnicas de fechamento que permitam aos vendedores concluir as negociações de forma positiva, garantindo que o cliente se sinta seguro e satisfeito com a decisão de compra.
        Reforce a importância de seguir os passos anteriores do método A PONTE para uma transição suave para o fechamento da venda, minimizando as objeções e maximizando as chances de sucesso.
        Exemplo Prático:

        "Se um cliente expressar preocupação com o custo, em vez de simplesmente oferecer um desconto, pergunte: 'Posso saber o que você considera como um investimento justo para este tipo de produto/serviço?' Use a resposta para destacar aspectos do seu produto que justifiquem o investimento, como qualidade superior, garantia estendida, ou suporte exclusivo."

        Detalhe a Etapa 'Negocie e Neutralize Objeções' do Método A PONTE, seguindo ESTRITAMENTE a estrutura abaixo. Cada seção deve ser claramente identificada e desenvolvida de forma extensa, mantendo o conteúdo genérico e aplicável a diversos setores:

        1. Objetivo:
        [Desenvolva este tópico detalhadamente]

        2. O que é negociar:
        [Desenvolva este tópico detalhadamente]

        3. Principais formas de pagamento:
        [Desenvolva este tópico detalhadamente, mantendo genérico]

        4. O que é neutralizar objeções:
        [Desenvolva este tópico detalhadamente]

        5. Como neutralizar objeções:
        [Desenvolva este tópico detalhadamente]

        6. Orientação financeira / venda consultiva:
        [Desenvolva este tópico detalhadamente]

        7. O que fazer:
        [Desenvolva este tópico detalhadamente]

        8. O que o vendedor ganha ao neutralizar as objeções?:
        [Desenvolva este tópico detalhadamente]

        9. O que o cliente ganha com a negociação?:
        [Desenvolva este tópico detalhadamente]

        10. Por que fazer?:
        [Desenvolva este tópico detalhadamente]

        11. Explique o que difere perguntas abertas e fechadas de negociação:
        [Desenvolva este tópico detalhadamente]

        12. Postura de negociação:
        [Desenvolva este tópico detalhadamente]

        13. Comportamento adequado diante das objeções:
        [Desenvolva este tópico detalhadamente]

        14. Pesquise e entenda as causas das objeções:
        [Desenvolva este tópico detalhadamente]

        15. Reforce os benefícios com foco na necessidade do cliente:
        [Desenvolva este tópico detalhadamente]

        16. Identificando as objeções:
        [Desenvolva este tópico detalhadamente]

        17. Catálogo de objeções comuns e respostas eficazes:
        [Desenvolva este tópico detalhadamente, incluindo exemplos genéricos aplicáveis a diversos setores]

        IMPORTANTE: 
        - Mantenha-se fiel a esta estrutura, desenvolvendo cada tópico separadamente e na ordem apresentada. 
        - Não adicione seções extras nem omita nenhuma das seções solicitadas.
        - Certifique-se de que cada seção seja claramente identificada com seu número e título correspondente.
        - Mantenha o conteúdo genérico, sem mencionar setores específicos ou exemplos de produtos particulares.
        - Evite redundâncias entre as seções, mas mantenha a coesão do texto como um todo.
        - Use perguntas abertas ao formular exemplos de respostas a objeções.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        
        """,

        "TOME A INICIATIVA E FECHE A VENDA": """

        Detalhe a Etapa 'Tome a Iniciativa e Feche a Venda' do Método A PONTE porém lembre-se de não ser redundante.

        Objetivo: Ensinar os vendedores a reconhecer o momento certo para fechar a venda, utilizando técnicas refinadas para garantir que o fechamento seja uma experiência positiva para o cliente, aumentando as chances de sucesso.

        Diretrizes:

        Reconhecimento de Sinais de Compra:

        Instrua os vendedores a ficarem atentos aos sinais verbais e não verbais de interesse do cliente, como perguntas sobre especificações do produto, comentários positivos, expressões faciais de satisfação ou gestos afirmativos.
        Enfatize a importância de interpretar esses sinais como uma abertura para iniciar o fechamento da venda.
        Técnicas de Fechamento:

        Apresente várias técnicas de fechamento e quando utilizá-las, como:
        Perguntas Alternativas: Encorajar o cliente a escolher entre opções, o que implica numa decisão final de compra.
        Supondo o Negócio Fechado: Falar como se o cliente já tivesse decidido comprar, levando-o a considerar como será após a compra.
        Fechamento Resumo: Resumir os benefícios e valor agregado do produto ou serviço discutidos durante a venda.
        Fechamento por Concessão: Oferecer algo adicional, como um desconto ou um bônus, para incentivar o cliente a fechar imediatamente.
        Evitar Táticas Pressionadoras:

        Advertir contra o uso de técnicas agressivas ou ultrapassadas que podem alienar o cliente. Em vez disso, focar em métodos que construam relacionamentos e confiança.
        Treine os vendedores a usar um tom de voz calmo e confiante, que tranquilize o cliente sobre sua decisão.
        Prática e Feedback:

        Realize sessões de role-playing para que os vendedores pratiquem as técnicas de fechamento em diferentes cenários de venda, recebendo feedback imediato para melhorar suas abordagens.
        Incentive os vendedores a refletirem sobre suas experiências de venda, identificando o que funcionou ou não e ajustando suas técnicas conforme necessário.
        

        Detalhe a Etapa 'Tome a Iniciativa e Feche a Venda' do Método A PONTE, seguindo ESTRITAMENTE a estrutura abaixo. Cada seção deve ser claramente identificada e desenvolvida de forma extensa, mantendo o conteúdo genérico e aplicável a diversos setores:

        1. Objetivo:
        [Desenvolva este tópico detalhadamente]

        2. Postura do parceiro para o fechamento do negócio:
        [Desenvolva este tópico detalhadamente]

        3. Gestão das informações do cliente:
        [Desenvolva este tópico detalhadamente]

        4. Sinais de compra do cliente:
        [Desenvolva este tópico detalhadamente]

        5. Técnicas de fechamento:
        [Desenvolva este tópico detalhadamente, incluindo várias técnicas aplicáveis a diferentes situações]

        6. O que o consultor ganha ao tomar a iniciativa e fechar a venda?:
        [Desenvolva este tópico detalhadamente]

        7. O que o cliente ganha?:
        [Desenvolva este tópico detalhadamente]

        8. Por que fazer?:
        [Desenvolva este tópico detalhadamente]

        9. Entenda as Decisões de Fechamento:
        [Desenvolva este tópico detalhadamente]

        10. Decisões Sempre Envolvem Riscos:
        [Desenvolva este tópico detalhadamente]

        11. ESTEJA ATENTO AOS SINAIS DE INTERESSE DO CLIENTE:
            a) Sinais verbais:
            [Desenvolva este subtópico detalhadamente]
            b) Sinais não verbais:
            [Desenvolva este subtópico detalhadamente]

        12. Use as principais técnicas de fechamento:
        [Desenvolva este tópico detalhadamente, expandindo sobre as técnicas mencionadas no item 5]

        IMPORTANTE: 
        - Mantenha-se fiel a esta estrutura, desenvolvendo cada tópico separadamente e na ordem apresentada. 
        - Não adicione seções extras nem omita nenhuma das seções solicitadas.
        - Certifique-se de que cada seção seja claramente identificada com seu número e título correspondente.
        - Mantenha o conteúdo genérico, sem mencionar setores específicos ou exemplos de produtos particulares.
        - Evite redundâncias entre as seções, mas mantenha a coesão do texto como um todo.
        - Forneça exemplos práticos e genéricos para ilustrar os conceitos, quando apropriado.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        
        """,

        "ESTENDA O RELACIONAMENTO": """

        Detalhe a Etapa 'ESTENDA O RELACIONAMENTO' do Método A PONTE porém lembre-se de não ser redundante.

        Objetivo: Ensinar os vendedores a fortalecer e prolongar o relacionamento com clientes, tanto aqueles que fecharam negócios quanto aqueles que não fecharam, utilizando estratégias eficazes de acompanhamento e engajamento contínuo.

        Diretrizes:

        Gestão de Informações do Cliente:

        Instrua os vendedores a coletar e registrar meticulosamente informações de contato e detalhes de interesse dos clientes durante e após as interações de venda, facilitando o acompanhamento futuro.
        Ressalte a importância de manter esses dados organizados e acessíveis, para que possam ser facilmente usados em campanhas de marketing ou follow-ups personalizados.
        Resgate de Clientes Não Convertidos:

        Treine os vendedores sobre como reengajar clientes que inicialmente não fecharam negócios, abordando suas objeções anteriores de maneira direcionada e oferecendo novas informações ou ajustes nas ofertas que possam atender melhor às suas necessidades.
        Encoraje a prática de entrar em contato rapidamente, preferencialmente dentro de 24 horas após a interação inicial, para manter a relevância e o interesse do cliente.
        
        Acompanhamento Pós-Venda:

        Ensine técnicas eficazes para o acompanhamento pós-venda, como enviar agradecimentos personalizados, solicitar feedback sobre a experiência de compra, e informar sobre novos produtos ou serviços que possam interessar ao cliente.
        Discuta estratégias para manter o relacionamento vivo, como contatos regulares em datas comemorativas, aniversários de compra, ou quando as parcelas de um produto estão se concluindo, incentivando novas compras.
        Pedidos de Indicação:

        Explique como e quando pedir indicações de maneira tática, transformando clientes satisfeitos em promotores da marca, que podem ajudar a expandir a base de clientes através de suas redes de contatos.
        Treinamento e Simulação:

        Realize sessões de treinamento que simulem diferentes cenários de follow-up com clientes, desde agradecimentos até negociações para resgatar um cliente não convertido.
        Utilize feedback dessas sessões para melhorar as abordagens e personalizar as interações de acordo com diferentes perfis de clientes.
        Exemplo Prático:

        "Após concluir uma venda, um vendedor pode enviar um e-mail de agradecimento ao cliente, incluindo detalhes de contato para suporte futuro e informações sobre produtos complementares que possam ser de interesse. Em seguida, marque um lembrete para entrar em contato novamente em um mês, oferecendo um desconto exclusivo para uma próxima compra ou solicitando uma avaliação do produto adquirido."

        Detalhe a Etapa 'Estenda o Relacionamento' do Método A PONTE, seguindo ESTRITAMENTE a estrutura abaixo. Cada seção deve ser claramente identificada e desenvolvida de forma extensa, mantendo o conteúdo genérico e aplicável a diversos setores:

        1. Objetivo:
        [Desenvolva este tópico detalhadamente]

        2. O que fazer com clientes que fecharam o negócio:
        [Desenvolva este tópico detalhadamente]

        3. O que fazer com clientes que não fecharam o negócio:
        [Desenvolva este tópico detalhadamente]

        4. Gerando mais negócios ao estender o relacionamento:
        [Desenvolva este tópico detalhadamente]

        5. Fluxo para indicações:
        [Desenvolva este tópico detalhadamente]

        6. Fluxo para novo contato com cliente:
        [Desenvolva este tópico detalhadamente]

        7. O que o consultor ganha ao estender o relacionamento?:
        [Desenvolva este tópico detalhadamente]

        8. O que o cliente ganha com o relacionamento estendido?:
        [Desenvolva este tópico detalhadamente]

        9. Por que fazer?:
        [Desenvolva este tópico detalhadamente]

        10. Acompanhe o cliente:
        [Desenvolva este tópico detalhadamente]

        11. Parabenize o cliente pela compra:
        [Desenvolva este tópico detalhadamente]

        12. Acompanhamento da venda:
        [Desenvolva este tópico detalhadamente]

        13. FOLLOW-UP:
            a) Importância do follow-up no processo de venda
            b) Melhores práticas para realizar follow-up eficaz
            c) Cronograma sugerido para diferentes tipos de follow-up
            d) Como usar o follow-up para superar objeções e fechar vendas
            e) Ferramentas e tecnologias para automatizar e melhorar o processo de follow-up
            f) Modelos de e-mails e scripts para diferentes situações de follow-up
        

        IMPORTANTE: 
        - Mantenha-se fiel a esta estrutura, desenvolvendo cada tópico separadamente e na ordem apresentada. 
        - Não adicione seções extras nem omita nenhuma das seções solicitadas.
        - Certifique-se de que cada seção seja claramente identificada com seu número e título correspondente.
        - Mantenha o conteúdo genérico, sem mencionar setores específicos ou exemplos de produtos particulares.
        - Evite redundâncias entre as seções, mas mantenha a coesão do texto como um todo.
        - Forneça exemplos práticos e genéricos para ilustrar os conceitos, quando apropriado.
        - Na seção de FOLLOW-UP, inclua modelos de e-mails e scripts genéricos que possam ser adaptados para diferentes situações e setores.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        
        """,
        
        "Conclusão": """
        Crie uma conclusão impactante e um plano de ação para o método APONTE no contexto de vendas de (segmento da empresa):

    
        1. Importância da implementação consistente:
        [Explique por que é crucial implementar o método APONTE de forma consistente nas vendas de (segmento da empresa)]

        
        2. Plano de ação (30, 60, 90 dias):
        [Detalhe um plano de ação para implementar o método APONTE em uma empresa de (segmento da empresa)]

        
        3. Mensagem motivacional final:
        [Inclua uma mensagem motivacional focada em vendas de (segmento da empresa)]

        Inclua citações motivacionais como:
        "O sucesso em vendas depende do que você faz para merecê-lo."
        "Intenção sem ação é ilusão: o mundo é de quem faz."

        Inclua um checklist de implementação e metas de curto e longo prazo.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        """,

        "CADERNO DE OBJEÇÕES": """
        Objetivo: Capacitar os vendedores a utilizar a ferramenta Caderno de Objeções para mapear objeções comuns e desenvolver respostas estratégicas que ajudem a neutralizar essas objeções, aumentando as chances de sucesso em vendas.

        Diretrizes:

        Entendimento das Objeções:

        Explicar que objeções são oportunidades de entender e atender melhor as necessidades do cliente.
        Ressaltar que cada "não" do cliente é um passo para um "sim" potencial, se abordado com a técnica correta.
        Técnicas de Neutralização:

        Demonstrar empatia: Mostrar ao cliente que você entende e respeita sua perspectiva.
        Devolver a objeção em forma de pergunta: Transformar a objeção do cliente em uma pergunta para entender melhor suas preocupações.
        Argumentar focando em benefícios: Responder com informações que destaquem o valor e os benefícios do produto ou serviço oferecido.
        Preenchimento do Caderno de Objeções:

        Instruir sobre como preencher o cabeçalho com nome, local de atuação e data.
        Orientar sobre como registrar uma objeção na coluna da esquerda e a estratégia de neutralização correspondente na coluna da direita.
        Detalhar como usar a seção de argumentação para desenvolver uma resposta focada nos benefícios do produto ou serviço.
        Prática e Revisão:

        Encorajar a prática constante, registrando e revisando as objeções e as respostas para refinamento contínuo.
        Sugerir compartilhar o Caderno de Objeções com colegas para feedback e novas ideias.
        Fechamento e Recurso Adicional:

        Oferecer o download da ferramenta Caderno de Objeções e encorajar seu uso frequente.
        Finalizar com uma chamada à ação para começar a jornada de melhoria na neutralização de objeções.
        Exemplo Prático:

        "Imagine que um cliente diz que o produto está caro. No Caderno de Objeções, você pode anotar essa objeção e, como resposta, perguntar: 'Entendo, está caro em relação a que especificamente?' Depois, dependendo da resposta do cliente, argumente sobre a qualidade superior e os benefícios exclusivos do seu produto."
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        """,

        "CADERNO DE BENEFÍCIOS": """
        Objetivo: Ensinar os vendedores a distinguir entre características e benefícios dos produtos ou serviços e a aplicar esse conhecimento de maneira estratégica para atender a perfis específicos de clientes, melhorando a eficácia das vendas.

        Diretrizes:

        Compreensão de Características vs. Benefícios:

        Explicar que características descrevem o produto ou serviço, enquanto benefícios focam no que o cliente ganha ao utilizar o produto ou serviço.
        Enfatizar que as vendas eficazes são guiadas pelos benefícios que resolvem problemas ou satisfazem necessidades específicas dos clientes.
        Identificação de Benefícios:

        Ensinar a listar as características de cada produto ou serviço.
        Orientar sobre como transformar cada característica em um benefício palpável para o cliente, perguntando "o que o cliente ganha com isso?".
        Aplicação Prática ao Perfil do Cliente:

        Instruir sobre como adaptar os benefícios de acordo com os perfis dos clientes, considerando fatores como idade, profissão, necessidades, e potencial financeiro.
        Usar exemplos práticos, como a diferenciação entre um jovem universitário e uma senhora aposentada, para mostrar como os benefícios devem ser apresentados de maneira diferente.
        Preenchimento do Caderno de Benefícios:

        Detalhar como preencher o cabeçalho com informações sobre o vendedor, data, e detalhes do produto ou serviço.
        Explicar como documentar a análise de benefícios e as abordagens personalizadas baseadas nos perfis de clientes em uma tabela tripartida: Características, Benefícios Gerais, e Benefícios Específicos para o Perfil.
        Prática e Revisão:

        Encorajar a prática diária inicialmente, reduzindo gradualmente à medida que os vendedores se tornam proficientes.
        Sugerir o preenchimento do caderno após cada atendimento, refletindo sobre o perfil do cliente e os benefícios mais relevantes discutidos.
        Fechamento e Recurso Adicional:

        Oferecer o download da ferramenta e recomendar seu uso regular como parte da preparação e revisão contínua.
        Incentivar os vendedores a compartilhar suas versões do caderno com colegas para feedback e aperfeiçoamento colaborativo.
        Exemplo Prático:

        "Suponha que você está vendendo um smartphone. Para um profissional que viaja frequentemente, enfatize o benefício da longa duração da bateria e a conectividade global. Para um estudante, destaque a câmera de alta qualidade para capturar momentos com amigos e a capacidade de gerenciar tarefas escolares eficientemente."
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        """,

        "RAIO-X": """
        Objetivo: Capacitar os vendedores a utilizar a ferramenta Raio X para avaliar a eficácia do seu atendimento de vendas, com base na metodologia A PONTE, e identificar áreas de melhoria.

        Diretrizes:

        Introdução à Ferramenta Raio X:

        Explique o propósito da ferramenta Raio X: avaliar a qualidade do atendimento de vendas e transformar essa avaliação em um indicador quantificável que impacta diretamente os resultados financeiros.
        Destaque que, embora a ferramenta seja exemplificada com a metodologia A PONTE, ela pode ser adaptada para outras metodologias de vendas.
        
        Funcionamento da Ferramenta:

        Descreva o Raio X como um checklist detalhado que acompanha as etapas da venda, com subetapas específicas que devem ser marcadas durante ou após um atendimento.
        Instrua sobre a utilização do checklist, seja por um líder ou pelo próprio vendedor como forma de autoavaliação.
        Etapas e Subetapas:

        Liste as etapas da metodologia A PONTE, incluindo suas subetapas e exemplos práticos de como marcar cada item durante o atendimento.
        Explique a pontuação: cada etapa vale até 10 pontos, com uma pontuação total possível de 60 pontos.
        Avaliação e Melhoria Contínua:

        Detalhe como interpretar as pontuações obtidas, correlacionando-as com a qualidade do atendimento (de Crítico a Excelente).
        Discuta a importância de realizar múltiplas avaliações para identificar padrões e áreas que necessitam de treinamento e melhoria.
        Fechamento:

        Encoraje o uso contínuo da ferramenta para melhorar o desempenho dos vendedores.
        Finalize com uma chamada para a ação, incentivando os vendedores a experimentar a ferramenta em seus próximos atendimentos.
        Exemplo Prático:

        "Ao abordar um cliente, verifique se o vendedor foi ao encontro do cliente, sorriu, desejou boas-vindas e se apresentou corretamente. Marque 'Sim' ou 'Não' para cada subetapa no Raio X."
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        """,
        
        "DE VOLTA À LOJA": """
        Crie uma ferramenta chamada "De Volta à Loja" relacionada à etapa E (Estenda o Relacionamento) do método APONTE. Esta ferramenta deve ser projetada para ajudar os vendedores a manter um relacionamento contínuo com os clientes e incentivá-los a retornar à loja.

        Objetivo da ferramenta:
        Fornecer estratégias e táticas práticas para que os vendedores mantenham contato com os clientes após a venda inicial, incentivando-os a voltar à loja para novas compras ou recomendações.

        Elementos-chave a serem incluídos:

        1. Título: "De Volta à Loja: Estratégias para Fidelização e Retorno de Clientes"

        2. Introdução:
        - Breve explicação da importância de manter o relacionamento com o cliente após a venda
        - Conexão com a etapa E do método APONTE

        3. Estratégias de acompanhamento pós-venda:
        - Definir um cronograma de contatos (por exemplo, 1 semana, 1 mês, 3 meses após a compra)
        - Sugestões de meios de contato (telefone, e-mail, mensagem de texto, redes sociais)
        - Exemplos de scripts ou modelos de mensagens para cada ponto de contato

        4. Programas de fidelidade:
        - Ideias para criar e implementar um programa de fidelidade
        - Benefícios e recompensas que podem ser oferecidos aos clientes recorrentes

        5. Eventos e promoções especiais:
        - Sugestões para eventos exclusivos para clientes existentes
        - Ideias de promoções sazonais ou personalizadas

        6. Coleta e uso de feedback:
        - Métodos para solicitar feedback dos clientes
        - Como usar o feedback para melhorar o serviço e fortalecer o relacionamento

        7. Técnicas de venda cruzada e upselling:
        - Identificação de oportunidades para oferecer produtos complementares
        - Estratégias para apresentar upgrades ou novos produtos relacionados

        8. Gerenciamento de informações do cliente:
        - Dicas para manter um registro atualizado das interações e preferências do cliente
        - Como usar essas informações para personalizar futuras interações

        9. Medição de sucesso:
        - KPIs para avaliar a eficácia das estratégias de retenção de clientes
        - Métodos para rastrear e analisar a taxa de retorno dos clientes

        10. Exemplos práticos:
            - Casos de estudo ou cenários hipotéticos demonstrando o uso bem-sucedido da ferramenta

        11. Checklist do vendedor:
            - Lista de verificação para garantir que todas as etapas da ferramenta "De Volta à Loja" sejam seguidas

        Formato:
        Apresente a ferramenta em um formato estruturado e fácil de ler, usando subtítulos, listas com marcadores e, se possível, sugestões de elementos visuais (como ícones ou diagramas) que poderiam ser incluídos para melhorar a compreensão e aplicação da ferramenta.

        Adaptação:
        Certifique-se de adaptar o conteúdo ao contexto de {opcoes_personalizacao['seguimento']} e às necessidades específicas da {opcoes_personalizacao['cliente_nome']}, considerando o tipo de produtos/serviços oferecidos e o perfil dos clientes.

        Lembre-se de manter a linguagem alinhada com o tom {', '.join(opcoes_personalizacao['tipos_linguagem'])} solicitado pelo cliente.

        Por favor, gere o conteúdo detalhado para esta ferramenta "De Volta à Loja", seguindo as diretrizes acima e mantendo-a alinhada com o método APONTE e as necessidades específicas do cliente.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        """,
        "MANUAL DE BOLSO": """
        Crie um Manual de Bolso prático e conciso para o vendedor, baseado na metodologia A PONTE. Este manual deve ser uma ferramenta de referência rápida que o vendedor possa consultar facilmente durante o processo de venda.

        Diretrizes:

        1. Formato conciso: O manual deve ser curto e direto, adequado para consulta rápida.
        2. Estrutura clara: Use tópicos, subtópicos e frases curtas para fácil leitura.
        3. Foco prático: Inclua dicas acionáveis e lembretes para cada etapa do método A PONTE.
        4. Personalização: Adapte o conteúdo ao segmento específico da empresa ({opcoes_personalizacao['seguimento']}).
        5. Exemplos rápidos: Forneça exemplos curtos e relevantes para cada etapa.
        6. Palavras-chave: Destaque palavras-chave importantes para cada etapa.
        7. Checklist: Inclua um checklist rápido para cada etapa do método A PONTE.
        8. Dicas de emergência: Adicione uma seção com dicas para lidar com situações difíceis ou objeções comuns.

        Estrutura sugerida:

        1. Introdução rápida ao método A PONTE
        2. Seções para cada etapa do A PONTE:
        - A - Aborde Positivamente
        - P - Pesquise o Cliente
        - O - Ofereça Soluções
        - N - Negocie e Neutralize Objeções
        - T - Tome a Iniciativa e Feche
        - E - Estenda o Relacionamento
        3. Checklist final
        4. Dicas de emergência

        Lembre-se de manter o conteúdo alinhado com as características específicas da {opcoes_personalizacao['cliente_nome']} e do segmento de {opcoes_personalizacao['seguimento']}.

        Forneça o conteúdo em formato Markdown, garantindo uma estrutura clara e fácil de ler.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        """
        }

        prompt_base = f"""
        Contexto: Você está criando um método de vendas personalizado para a {opcoes_personalizacao['cliente_nome']}, uma empresa do segmento de {opcoes_personalizacao['seguimento']} - {opcoes_personalizacao['seguimento_desc']}. O método APONTE é a base deste treinamento, mas deve ser adaptado às necessidades específicas da {opcoes_personalizacao['cliente_nome']} e do mercado de {opcoes_personalizacao['seguimento']}.

        Perfil do Consultor de Sucesso em Vendas:
        Objetivo Principal: Transformar pessoas e empresas por meio de métodos comprovados de vendas, alinhando performance humana, tecnologia e processos para obter resultados excepcionais.
        Valores Centrais:
        Inovação na prática: Implementar soluções tangíveis e criativas que impulsionem negócios e pessoas.
        

        Abordagem e Metodologia

            1.	Ensinar com base na prática e evidência:
            •	Vendas é vista como um processo técnico e metódico, que pode ser aprendido, aplicado e melhorado continuamente    .
            •	O método “A PONTE” é um exemplo de abordagem estruturada para guiar o cliente através do funil de vendas  .
            2.	Liderança Treinadora:
            •	Envolve-se ativamente no desenvolvimento da equipe, focando em motivação, engajamento e habilidade prática.
            •	Oferece treinamento contínuo baseado em resultados reais e customizados    .
            3.	Foco em Resultados:
            •	Reconhece que o resultado é consequência de ações bem estruturadas e de hábitos consistentes.
            •	A gestão de resultados vai além da análise numérica, conectando números com ações práticas e feedback contínuo    .

        Personalidade e Tom de Voz

            •	Transformador e Inspirador:
            •	Incentiva a evolução das pessoas, conectando o sucesso individual ao coletivo  .
            •	Prático e Estratégico:
            •	Utiliza dados, tecnologia e metodologias comprovadas, mas foca em ações que criam impacto direto    .
            •	Empático e Proativo:
            •	Fala de forma próxima, humana e inspiradora, promovendo colaboração e parcerias com clientes e equipes    .

        Ferramentas e Competências

            •	Conhecimento técnico e prático:
            •	Metodologia robusta: APONTE
        

            Crie a seção "{parte}" do método de vendas personalizado para o cliente, incorporando o método APONTE.


        IMPORTANTE:
        1. Siga ESTRITAMENTE a estrutura fornecida para a seção "{parte}".
        2. Cada tópico deve ser claramente identificado pelo seu título.
        3. Desenvolva cada tópico detalhadamente, mantendo o foco no segmento de {opcoes_personalizacao['seguimento']}.
        4. Use exemplos práticos e casos de uso específicos para a {opcoes_personalizacao['cliente_nome']}.
        5. Mantenha um tom {', '.join(opcoes_personalizacao['tipos_linguagem'])}.
        6. Foque nos seguintes canais de atendimento: {', '.join(opcoes_personalizacao['canais_atendimento'])}.
        7. {gerar_instrucoes_tamanho(opcoes_personalizacao['tamanho_material'])}
        8. {gerar_instrucoes_segmento(opcoes_personalizacao['seguimento'])}
        9. Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal. Evite muitos bullet points, traga também parágrafos explicativos e discursivos.


        Lembre-se que {opcoes_personalizacao['cliente_nome']} é o grupo, qual o material está sendo construido.
        Lembre-se de manter um fluxo lógico e evitar redundâncias entre as seções.

        {prompts.get(parte, "")}
        Use as seguintes informações como referência:
        Conteúdo dos arquivos: {conteudo_arquivos[:1000000]}...
        Materiais: {materiais[:1000000]}...
        Modelo base: {modelo_metodo_vendas[:1000000]}...

        Forneça o conteúdo em formato Markdown, garantindo uma estrutura clara e fácil de ler.

        """

        mensagem = HumanMessage(content=prompt_base)
        resposta = modelo_ia.invoke([mensagem])
        parte_metodo = resposta.content if hasattr(resposta, 'content') else str(resposta)
        
        def correcao_o_que_e_o_metodo(parte_metodo, opcoes_personalizacao):
            prompt_revisao = f"""
            Revise e reescreva o seguinte conteúdo sobre "O que é o Método de Vendas A PONTE":

            {parte_metodo}

            Instruções específicas:

            1. Forneça um resumo conciso e claro sobre o que é o método APONTE.
            2. Não liste as etapas do método em tópicos. Em vez disso, integre-as em um texto fluido.
            3. Explique brevemente a importância e o propósito geral do método.
            4. Mencione que cada etapa será detalhada nas seções seguintes do material.
            5. Mantenha o foco no segmento de {opcoes_personalizacao['seguimento']} e nas necessidades da {opcoes_personalizacao['cliente_nome']}.
            6. Use uma linguagem {', '.join(opcoes_personalizacao['tipos_linguagem'])}.
            7. Trazer tabela, nome da etapa do metodo APONTE e curta definição.

            Forneça um conteúdo conciso, informativo e envolvente que introduza o método sem entrar em detalhes específicos de cada etapa.
            """

            mensagem_revisao = HumanMessage(content=prompt_revisao)
            resposta_revisao = modelo_ia.invoke([mensagem_revisao])
            return resposta_revisao.content if hasattr(resposta_revisao, 'content') else str(resposta_revisao)

        # Função de correção geral para outras partes
        def correcao_geral(parte, parte_metodo, opcoes_personalizacao):
            prompt_revisao = f"""
            Revise e reescreva o seguinte conteúdo gerado para a seção "{parte}" do método de vendas:

            {parte_metodo}

            Instruções de revisão e reescrita:

            1. Reescreva todo o conteúdo, mesmo que não haja necessidade de correções significativas.
            2. Elimine todas as redundâncias e informações repetitivas.
            3. Expanda e elabore todos os pontos, fornecendo mais detalhes, exemplos e explicações.
            4. Não use frases como "Manter como no original" ou "(Manter ... como no original)". Em vez disso, reescreva completamente essas seções.
            5. Para tabelas ou listas mencionadas no texto original, recrie-as com informações detalhadas e exemplos adicionais.
            6. Para exemplos de abordagem, forneça novos exemplos específicos para cada canal de atendimento, focando em diferentes produtos reais e situações.
            7. Mantenha o foco no segmento de {opcoes_personalizacao['seguimento']} e nas necessidades específicas da {opcoes_personalizacao['cliente_nome']}.
            8. Garanta que todo o conteúdo seja relevante e contribua para o método de vendas APONTE.
            9. Use uma linguagem clara, direta e envolvente, adequada ao tom {', '.join(opcoes_personalizacao['tipos_linguagem'])}.
            10. Forneça um conteúdo completo e detalhado, sem omissões ou referências a um conteúdo anterior não incluído.

            Lembre-se: Seu objetivo é fornecer um conteúdo totalmente reescrito, expandido e melhorado, sem fazer referência a versões ou conteúdos anteriores.
            """

            mensagem_revisao = HumanMessage(content=prompt_revisao)
            resposta_revisao = modelo_ia.invoke([mensagem_revisao])
            return resposta_revisao.content if hasattr(resposta_revisao, 'content') else str(resposta_revisao)

        # Escolha a função de correção apropriada
        if parte == "O que é o Método de Vendas A PONTE":
            parte_metodo_revisada = correcao_o_que_e_o_metodo(parte_metodo, opcoes_personalizacao)
        else:
            parte_metodo_revisada = correcao_geral(parte, parte_metodo, opcoes_personalizacao)

        if parte_metodo_revisada != parte_metodo:
            logger.info(f"O conteúdo da seção '{parte}' foi completamente reescrito e expandido.")

        return parte_metodo_revisada

    # Função para gerar anexos
    def gerar_anexos(conteudo_arquivos, materiais):
        prompt_anexos = f"""
        Crie os seguintes anexos para o método de vendas do cliente:
        
        A - Catálogo de Benefícios
        B - Catálogo de Objeções
        C - Catálogo de Produtos
        D - Tabela Comparativa de Produtos
        
        Para cada anexo:
        1. Foque em informações únicas e relevantes para o cliente.
        2. Evite repetir informações entre os anexos.
        3. Use dados específicos extraídos dos materiais fornecidos.
        4. Organize as informações de forma lógica e fácil de consultar.
        5. Para o Catálogo de Objeções, inclua respostas concisas e eficazes.
        6. Na Tabela Comparativa, destaque as vantagens competitivas dos produtos do cliente.
        
        Use o formato Markdown e crie tabelas quando apropriado para melhor visualização.
        Não use blocos de código ou formatação especial. Apresente todo o conteúdo como texto normal, usando tabelas quando necessario.
        
        Informações disponíveis:
        Conteúdo dos arquivos: {conteudo_arquivos[:1000000]}...
        Materiais: {materiais[:1000000]}...
        """
        
        mensagem = HumanMessage(content=prompt_anexos)
        resposta = modelo_ia.invoke([mensagem])
        anexos = resposta.content if hasattr(resposta, 'content') else str(resposta)
        
        return anexos

    # Função principal para gerar o método de vendas
    def gerar_metodo_vendas(conteudo_arquivos, materiais, modelo_metodo_vendas, opcoes_personalizacao):
        st.write("Iniciando a geração do método de vendas...")
        metodo_completo = gerar_metodo_vendas_completo(conteudo_arquivos, materiais, modelo_metodo_vendas, opcoes_personalizacao)
        
        st.write("Gerando anexos...")
        anexos = gerar_anexos(conteudo_arquivos, materiais)
        
        metodo_final = f"{metodo_completo}\n\n# Anexos\n\n{anexos}"
        return metodo_final

    def salvar_metodo_vendas(conteudo):
        html = markdown2.markdown(conteudo, extras=["tables", "fenced-code-blocks"])
        soup = BeautifulSoup(html, 'html.parser')

        doc = Document()
        
        styles = doc.styles
        style_normal = styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(11)

        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table', 'pre']):
            if element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                doc.add_heading(element.text, level=level)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet' if element.name == 'ul' else 'List Number')
            elif element.name == 'table':
                rows = element.find_all('tr')
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.text.strip()
            elif element.name == 'pre':
                # Tratar blocos de código como parágrafos normais
                doc.add_paragraph(element.text)

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return tmp.name

    def gerar_conteudo_email(nome_consultor, cliente_nome, opcoes_personalizacao, arquivos_utilizados):
        prompt = f"""
        Crie um e-mail personalizado para enviar o método de vendas gerado o líder do projeto que é Danielle, para o diretor de operações e consultoria Marcio Godoy, e para o PMO da empresa Emanuel.

        - Consultor: {nome_consultor}
        - Cliente: {cliente_nome}
        - Data e hora de geração: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        - Opções de personalização: {json.dumps(opcoes_personalizacao, indent=2)}
        - Arquivos utilizados: {', '.join(arquivos_utilizados)}

        O e-mail deve incluir:
        
        1. Uma breve introdução explicando o propósito do e-mail
        2. Um resumo das principais características do método de vendas gerado
        3. Uma seção destacando as opções de personalização usadas
        4. Uma lista dos arquivos utilizados na geração do método
        
        Use formatação HTML para melhorar a apresentação do e-mail, incluindo negrito, itálico, e listas onde apropriado.
        """

        mensagem = HumanMessage(content=prompt)
        resposta = modelo_ia.invoke([mensagem])
        return resposta.content if hasattr(resposta, 'content') else str(resposta)

    def enviar_email(destinatarios, assunto, corpo, anexos):
        remetente = os.getenv('EMAIL_REMETENTE', "gerador.ia.sv@gmail.com")
        senha = os.getenv('EMAIL_SENHA', "adhjcqqqthcrrdpo")

        if not remetente or not senha:
            logger.error("Credenciais de e-mail não encontradas nas variáveis de ambiente.")
            return False

        try:
            yag = yagmail.SMTP(remetente, senha)
            
            contents = [corpo]  # O corpo do e-mail agora é HTML gerado pelo Gemini
            for anexo in anexos:
                if os.path.exists(anexo):
                    contents.append(anexo)
                else:
                    logger.warning(f"Anexo não encontrado: {anexo}")

            yag.send(to=destinatarios, subject=assunto, contents=contents)
            
            logger.info(f"E-mail enviado com sucesso para {', '.join(destinatarios)}")
            return True
        except Exception as e:
            logger.error(f"Erro ao enviar e-mail: {str(e)}")
            return False

    # CSS básico para estilo consistente
    css = """
    <style>
        .centered-title {
            text-align: center;
            padding: 20px 0;
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

    # Interface do usuário Streamlit
    col1, col2, col3 = st.columns([1,2,1])

    with col2:
        try:
            logo_path = assets_dir / "LOGO SUCESSO EM VENDAS HORIZONTAL AZUL.png"
            st.image(str(logo_path), width=300, use_container_width=True)
        except FileNotFoundError:
            st.write("Logo não encontrada. Por favor, verifique o caminho da imagem.")

    st.markdown("<h1 class='centered-title'>Gerador de Método de Vendas Varejo</h1>", unsafe_allow_html=True)
    st.write("")

    # Inicializar o estado da sessão
    if 'materiais_adicionais' not in st.session_state:
        st.session_state.materiais_adicionais = []

    # Carregar o modelo de método de vendas
    modelo_metodo_vendas = ""
    modelo_path = Path(__file__).parent.parent / "Metodo_de_vendas.docx"
    try:
        with open(modelo_path, "rb") as arquivo:
            doc = Document(arquivo)
            modelo_metodo_vendas = "\n".join([p.text for p in doc.paragraphs])
    except FileNotFoundError:
        st.warning(f"Arquivo '{modelo_path}' não encontrado. Por favor, verifique se o arquivo existe no diretório correto.")
    except Exception as e:
        st.error(f"Erro ao carregar o modelo de método de vendas: {e}")

    # Barra lateral
    with st.sidebar:
        st.header("Informações do Consultor")
        nome_consultor = st.text_input("Nome do Consultor")
        
        st.header("Nome do Cliente")
        cliente_nome = st.text_input("Insira o nome do Cliente:", "")
        
        st.header("Gerador de Dossiê")
        url = st.text_input("URL do site da empresa")
        
        # Configurações de personalização
        seguimento = st.selectbox(
            "Selecione o seguimento da empresa",
            ["Concessionárias de veículos","Imóveis","Eletromóveis","Serviços Financeiros","Cama, Mesa e Banho","Tintas","Farma","Mat. Construção","Outros"], 
            index=1
        )
        seguimento_desc = st.text_input("Particularidades do Cliente", "(segmento da empresa)")

        destinatarios = st.multiselect(
            "Para quem se destina",
            ["Vendedor", "Gerente", "Supervisor", "Diretor"],
            default=["Vendedor"]
        )

        tipos_linguagem = st.multiselect(
            "Tipo de Linguagem",
            ["Formal", "Informal", "Técnico", "Consultivo", "Popular"],
            default=["Informal", "Popular"]
        )
        canais_atendimento = st.multiselect(
            "Canais de Atendimento",
            ["WhatsApp", "Telefone", "Presencial", "PaP"],
        )
        tamanho_material = st.selectbox(
            "Extensão do Material",
            ["Curto", "Médio", "Longo"],
            index=1
        )

        opcoes_personalizacao = {
            "tipos_linguagem": tipos_linguagem,
            "destinatarios": destinatarios,
            "canais_atendimento": canais_atendimento,
            "tamanho_material": tamanho_material,
            "seguimento": seguimento,
            "seguimento_desc": seguimento_desc,
            "cliente_nome": cliente_nome
        }

        st.header("Carregar Materiais")

        # Upload de arquivos
        entrevistas_vendedores = st.file_uploader("Carregar entrevistas com vendedores, gerente, supervisores e ou diretores", type=["txt", "docx", "pdf", "xlsx", "csv"], accept_multiple_files=True, key="vendedores")
        
        laboratorio_vendas = st.file_uploader("Carregar materiais do laboratório de vendas", type=["txt", "docx", "pdf", "xlsx", "csv"], accept_multiple_files=True, key="laboratorio")
        materiais_adicionais = st.file_uploader("Carregar materiais adicionais (catálogos, etc.)", type=["txt", "docx", "pdf", "xlsx", "csv", "json"], accept_multiple_files=True, key="adicionais")

        # Botão para gerar o método de vendas
        if st.button("Gerar Método de Vendas"):
            if not cliente_nome or not nome_consultor:
                st.error("Por favor, preencha todas as informações necessárias.")
            else:
                st.session_state.gerar_metodo = True
                
        # Botão para voltar à página inicial
        st.markdown("---")
        if st.button("← Voltar para a página inicial", key="btn_voltar_metodo"):
            st.session_state.current_app = 'home'
            st.query_params["app"] = "home"
            st.rerun()

    # Combinando todos os arquivos carregados
    arquivos_upload = []
    if entrevistas_vendedores:
        arquivos_upload.extend(entrevistas_vendedores)
    if laboratorio_vendas:
        arquivos_upload.extend(laboratorio_vendas)
    if materiais_adicionais:
        arquivos_upload.extend(materiais_adicionais)

    # Container principal
    if 'gerar_metodo' in st.session_state and st.session_state.gerar_metodo:
        # Gerar o dossiê primeiro se URL for fornecida
        caminho_dossie = None
        if url:
            with st.spinner("Gerando dossiê da empresa..."):
                caminho_dossie = gerar_e_salvar_dossie(url, cliente_nome)
                
                if caminho_dossie:
                    # Adicionar o dossiê aos materiais adicionais
                    with open(caminho_dossie, "rb") as dossie_file:
                        dossie_bytes = dossie_file.read()
                    
                    dossie_file_like = io.BytesIO(dossie_bytes)
                    dossie_file_like.name = os.path.basename(caminho_dossie)
                    
                    if 'materiais_adicionais' not in st.session_state:
                        st.session_state.materiais_adicionais = []
                    st.session_state.materiais_adicionais.append(dossie_file_like)
                    st.success("Dossiê gerado com sucesso e adicionado aos materiais!")
        
        # Continuar com a geração do método de vendas
        arquivos_upload += st.session_state.materiais_adicionais
        
        with st.spinner("Processando arquivos carregados..."):
            conteudo_arquivos = "\n\n".join([carregar_arquivo(arquivo) for arquivo in arquivos_upload])
            conteudo_materiais = carregar_conteudo_pasta(str(materials_dir))

        if not modelo_metodo_vendas:
            st.error("O modelo de método de vendas não foi carregado. Não é possível gerar o método personalizado.")
        else:
            with st.spinner("Gerando método de vendas personalizado..."):
                try:
                    logger.info(f"Tamanho do conteúdo dos arquivos carregados: {len(conteudo_arquivos)} caracteres")
                    logger.info(f"Tamanho do conteúdo dos materiais: {len(conteudo_materiais)} caracteres")

                    metodo_vendas = gerar_metodo_vendas(conteudo_arquivos, conteudo_materiais, modelo_metodo_vendas, opcoes_personalizacao)

                    st.markdown("## Método de Vendas Personalizado")
                    st.markdown(metodo_vendas)

                    try:
                        nome_metodo = f"Metodo_Vendas_{cliente_nome.replace(' ', '_')}.docx"
                        caminho_metodo = os.path.join(tempfile.gettempdir(), nome_metodo)
                        
                        # Modificar a função salvar_metodo_vendas para aceitar o caminho como argumento
                        caminho_docx = salvar_metodo_vendas(metodo_vendas)
                        
                        # Mover o arquivo gerado para o caminho desejado
                        shutil.move(caminho_docx, caminho_metodo)
                        
                        with open(caminho_metodo, "rb") as file:
                            docx_bytes = file.read()
                        
                        st.download_button(
                            label="Download DOCX",
                            data=docx_bytes,
                            file_name=nome_metodo,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        destinatarios = ["leonardo@sucessoemvendas.com.br"]
                        assunto = f"Método de Vendas Personalizado - {cliente_nome}"
                        
                        arquivos_utilizados = [arquivo.name for arquivo in arquivos_upload]
                        corpo_email = gerar_conteudo_email(nome_consultor, cliente_nome, opcoes_personalizacao, arquivos_utilizados)
                        
                        anexos = [caminho_metodo]
                        if caminho_dossie:
                            anexos.append(caminho_dossie)

                        with st.spinner("Enviando e-mail..."):
                            sucesso = enviar_email(destinatarios, assunto, corpo_email, anexos)
                        
                        if sucesso:
                            st.success("Método de Vendas gerado e enviado com sucesso!")
                        else:
                            st.error("Falha ao enviar o e-mail. Por favor, tente novamente.")

                        # Limpar arquivos temporários
                        try:
                            os.remove(caminho_metodo)
                            if caminho_dossie:
                                os.remove(caminho_dossie)
                        except:
                            pass
                            
                    except Exception as e:
                        st.error(f"Erro ao gerar arquivo DOCX ou enviar e-mail: {str(e)}")
                        logger.error(f"Erro ao gerar DOCX ou enviar e-mail: {e}", exc_info=True)

                except Exception as e:
                    st.error(f"Erro ao gerar o método de vendas: {str(e)}")
                    logger.error(f"Erro detalhado: {e}", exc_info=True)

        st.session_state.gerar_metodo = False

if __name__ == "__main__":
    app()