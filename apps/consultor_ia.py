import streamlit as st
import os
import json
import time
from datetime import datetime
import logging
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import ChatPromptTemplate
from langchain.globals import set_verbose
import tiktoken
import hashlib
import re
import docx
import PyPDF2
import pickle
import io
import sys
from pathlib import Path

# Adiciona o diretório raiz ao path para importar módulos corretamente
sys.path.append(str(Path(__file__).parent.parent))

def app(config=None):
    # Verifica se a página já foi configurada pelo app principal
    if not config or not config.get("already_configured"):
        # Configuração da página Streamlit (só será executada se o app for executado sozinho)
        st.set_page_config(page_title='Consultor I.A. Sucesso em Vendas', layout="wide")

    # Configuração de logging
    log_dir = Path(__file__).parent.parent / "logs"
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / "consultor_ia.log"
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(str(log_file)),
            logging.StreamHandler()
        ]
    )
    logger = logging.getLogger(__name__)

    # Definir a verbosidade
    set_verbose(True)

    # Caminho para os assets e materiais
    assets_dir = Path(__file__).parent.parent / "assets"
    materials_dir = Path(__file__).parent.parent / "materiais"
    materials_dir.mkdir(exist_ok=True)  # Cria a pasta se não existir
    data_dir = Path(__file__).parent.parent / "data"
    data_dir.mkdir(exist_ok=True)
    chats_file = data_dir / "consultor_chats.pkl"

    # Agora carrega o .env
    try:
        from dotenv import load_dotenv
        load_dotenv(override=True)  # Force override existing env variables
        logger.info(f"Novo caminho das credenciais: {os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
    except Exception as e:
        logger.error(f"Erro ao carregar .env: {e}")
        st.warning("Arquivo .env não encontrado. Algumas funcionalidades podem não estar disponíveis.")

    # Configuração do Gemini AI
    try:
        # Obter caminho das credenciais do .env
        caminho_credenciais = os.getenv('GOOGLE_APPLICATION_CREDENTIALS', 
                                       str(Path(__file__).parent.parent / "decent-atlas-460512-g7-3b1d4ccb9c4e.json"))
          # Verificar se o arquivo existe
        if not os.path.exists(caminho_credenciais):
            logger.warning(f"Arquivo de credenciais não encontrado: {caminho_credenciais}")
            st.warning("⚠️ Credenciais do Google Cloud não configuradas corretamente. Verifique a configuração.")
        # Configurar a variável de ambiente para o SDK do Google
        os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = caminho_credenciais
        
        # Inicializar o modelo
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.3)
        logger.info(f"Modelo Gemini inicializado com arquivo de credenciais: {caminho_credenciais}")
        
    except Exception as e:
        logger.error(f"Erro ao inicializar o modelo Gemini: {e}")
        st.error(f"Erro ao inicializar o modelo de IA: {str(e)}")
        llm = None

    # Função para contar tokens
    def num_tokens_from_string(string: str, model_name: str = "gpt-3.5-turbo") -> int:
        try:
            encoding = tiktoken.encoding_for_model(model_name)
            num_tokens = len(encoding.encode(string))
            return num_tokens
        except:
            # Fallback simples se tiktoken falhar
            return len(string.split())

    # Função para contar caracteres
    def count_characters(text):
        return len(text)

    # Função para carregar e processar arquivos JSON
    def load_json(file):
        try:
            if isinstance(file, (str, Path)):
                # Se for um caminho de arquivo
                with open(file, 'r') as f:
                    data = json.load(f)
            else:
                # Se for um objeto de arquivo (upload)
                data = json.load(file)
            return data
        except json.JSONDecodeError as e:
            raise ValueError(f"Erro ao decodificar o JSON: {e}")
        except Exception as e:
            raise RuntimeError(f"Erro ao carregar o arquivo JSON: {e}")

    # Função para carregar e processar arquivos DOCX
    def load_docx(file):
        try:
            if isinstance(file, (str, Path)):
                # Se for um caminho de arquivo
                doc = docx.Document(file)
            else:
                # Se for um objeto de arquivo (upload)
                doc = docx.Document(io.BytesIO(file.read()))
                file.seek(0)  # Resetar o ponteiro do arquivo
            
            text = "\n".join([p.text for p in doc.paragraphs])
            return text
        except Exception as e:
            raise RuntimeError(f"Erro ao carregar o arquivo DOCX: {e}")

    # Função para carregar e processar arquivos PDF
    def load_pdf(file):
        try:
            if isinstance(file, (str, Path)):
                # Se for um caminho de arquivo
                with open(file, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() if page.extract_text() else ""
            else:
                # Se for um objeto de arquivo (upload)
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() if page.extract_text() else ""
                file.seek(0)  # Resetar o ponteiro do arquivo
                
            return text
        except Exception as e:
            raise RuntimeError(f"Erro ao carregar o arquivo PDF: {e}")

    # Função para carregar todos os arquivos na pasta materiais
    def load_fixed_materials():
        materials = []
        total_tokens = 0
        total_chars = 0
        
        if not materials_dir.exists():
            logger.warning(f"Pasta de materiais não encontrada: {materials_dir}")
            return "", 0, 0

        for filename in os.listdir(materials_dir):
            filepath = materials_dir / filename
            try:
                if filename.endswith('.json'):
                    content = load_json(filepath)
                    content_str = str(content)
                    materials.append(content_str)
                    total_tokens += num_tokens_from_string(content_str)
                    total_chars += count_characters(content_str)
                    logger.info(f"Carregado material JSON: {filename}")
                elif filename.endswith('.docx'):
                    content = load_docx(filepath)
                    materials.append(content)
                    total_tokens += num_tokens_from_string(content)
                    total_chars += count_characters(content)
                    logger.info(f"Carregado material DOCX: {filename}")
                elif filename.endswith('.pdf'):
                    content = load_pdf(filepath)
                    materials.append(content)
                    total_tokens += num_tokens_from_string(content)
                    total_chars += count_characters(content)
                    logger.info(f"Carregado material PDF: {filename}")
            except Exception as e:
                logger.error(f"Erro ao carregar arquivo {filename}: {e}")
        
        materials_text = "\n\n".join(materials)
        logger.info(f"Total de tokens nos materiais fixos: {total_tokens}")
        logger.info(f"Total de caracteres nos materiais fixos: {total_chars}")
        return materials_text, total_tokens, total_chars

    # Função para processar arquivos carregados pelo usuário
    def process_uploaded_files(uploaded_files):
        if not uploaded_files:
            return "", 0, 0
            
        materials = []
        total_tokens = 0
        total_chars = 0
        
        for file in uploaded_files:
            try:
                if file.name.endswith('.json'):
                    content = load_json(file)
                    content_str = str(content)
                    materials.append(content_str)
                    total_tokens += num_tokens_from_string(content_str)
                    total_chars += count_characters(content_str)
                    logger.info(f"Carregado material JSON do usuário: {file.name}")
                elif file.name.endswith('.docx'):
                    content = load_docx(file)
                    materials.append(content)
                    total_tokens += num_tokens_from_string(content)
                    total_chars += count_characters(content)
                    logger.info(f"Carregado material DOCX do usuário: {file.name}")
                elif file.name.endswith('.pdf'):
                    content = load_pdf(file)
                    materials.append(content)
                    total_tokens += num_tokens_from_string(content)
                    total_chars += count_characters(content)
                    logger.info(f"Carregado material PDF do usuário: {file.name}")
                else:
                    st.warning(f"Formato de arquivo não suportado: {file.name}")
            except Exception as e:
                logger.error(f"Erro ao processar arquivo {file.name}: {e}")
                st.error(f"Erro ao processar arquivo {file.name}: {str(e)}")
        
        materials_text = "\n\n".join(materials)
        logger.info(f"Total de tokens nos materiais do usuário: {total_tokens}")
        logger.info(f"Total de caracteres nos materiais do usuário: {total_chars}")
        return materials_text, total_tokens, total_chars

    # Contexto fixo do agente
    agent_context = (
        "Você é um agente inteligente e consultor comercial da empresa Sucesso em Vendas. "
        "Gostaria que me respondesse de forma objetiva e concisa, com uma explicação sobre e em seguida uma abordagem pratica de como fazer para resolver. "
        "Seu papel é fornecer assistência especializada utilizando o método de vendas da Sucesso em Vendas e ajudar com conselhos comerciais para gerentes, coordenadores e vendedores."
    )

    # Função para gerar a resposta
    def generate_response(user_input, context):
        if not llm:
            return "Erro: Modelo de IA não inicializado. Verifique as credenciais."
            
        # Gerar uma chave única para o cache
        cache_key = hashlib.md5((user_input + context[:100]).encode()).hexdigest()
        
        # Verificar se a resposta está no cache
        if cache_key in st.session_state.response_cache:
            logger.info("Resposta encontrada no cache")
            cached_response = st.session_state.response_cache[cache_key]
            return cached_response

        prompt = f"{context}\n\nUsuário: {user_input}\nChatbot:"
        input_tokens = num_tokens_from_string(prompt)
        input_chars = count_characters(prompt)
        logger.info(f"Tokens na entrada: {input_tokens}")
        logger.info(f"Caracteres na entrada: {input_chars}")
        
        model = ChatPromptTemplate.from_template(prompt) | llm
        try:
            response = model.invoke({'input': prompt})
            response_content = response.content if hasattr(response, 'content') else str(response)
            
            response_tokens = num_tokens_from_string(response_content)
            response_chars = count_characters(response_content)
            logger.info(f"Tokens na resposta: {response_tokens}")
            logger.info(f"Caracteres na resposta: {response_chars}")
            
            total_tokens = input_tokens + response_tokens
            total_chars = input_chars + response_chars
            logger.info(f"Total de tokens nesta interação: {total_tokens}")
            logger.info(f"Total de caracteres nesta interação: {total_chars}")
            
            # Armazenar a resposta no cache
            st.session_state.response_cache[cache_key] = response_content
            
            return response_content
        except Exception as e:
            logger.error(f"Erro ao gerar resposta: {str(e)}")
            return f"Ocorreu um erro ao gerar a resposta: {str(e)}. Por favor, tente novamente."

    # Função para exibir a resposta gradualmente como se estivesse digitando
    def display_typing_response(response_text, container):
        typing_speed = 0.01  # Velocidade de digitação (em segundos por caractere)
        typed_text = ""
        for char in response_text:
            typed_text += char
            container.markdown(typed_text)
            time.sleep(typing_speed)

    # Função para extrair título do chat
    def extract_title(message):
        # Extrair as primeiras duas ou três palavras significativas
        words = re.findall(r'\b\w+\b', message)
        if len(words) >= 2:
            return f"{words[0]} {words[1]}..."
        return "Novo Chat"

    # Função para salvar os chats
    def save_chats():
        try:
            with open(chats_file, 'wb') as f:
                pickle.dump(st.session_state.chats, f)
            logger.info("Chats salvos com sucesso.")
        except Exception as e:
            logger.error(f"Erro ao salvar chats: {e}")
            st.error(f"Erro ao salvar chats: {str(e)}")

    # Função para carregar os chats
    def load_chats():
        if not chats_file.exists():
            logger.info("Arquivo de chats não encontrado. Criando novo.")
            return {'chat_1': {'date': datetime.now().strftime("%d/%m/%Y"), 
                               'messages': [], 
                               'title': "Novo Chat",
                               'user_materials': ""}}
        
        try:
            with open(chats_file, 'rb') as f:
                chats = pickle.load(f)
            logger.info(f"Chats carregados com sucesso. Total: {len(chats)}")
            return chats
        except Exception as e:
            logger.error(f"Erro ao carregar chats: {e}")
            st.error(f"Erro ao carregar chats: {str(e)}")
            return {'chat_1': {'date': datetime.now().strftime("%d/%m/%Y"), 
                               'messages': [], 
                               'title': "Novo Chat",
                               'user_materials': ""}}

    # Função para criar um novo chat
    def new_chat():
        current_date = datetime.now().strftime("%d/%m/%Y")
        chat_id = f"chat_{len(st.session_state.chats) + 1}"
        st.session_state.chats[chat_id] = {
            'date': current_date, 
            'messages': [], 
            'title': "Novo Chat",
            'user_materials': ""
        }
        st.session_state.current_chat_id = chat_id
        st.session_state.uploaded_files = []
        logger.info(f"Novo chat criado: {chat_id}")
        save_chats()

    # Função para renomear um chat
    def rename_chat(chat_id, new_title):
        if chat_id in st.session_state.chats:
            st.session_state.chats[chat_id]['title'] = new_title
            logger.info(f"Chat {chat_id} renomeado para: {new_title}")
            save_chats()

    # Função para excluir um chat
    def delete_chat(chat_id):
        if chat_id in st.session_state.chats:
            del st.session_state.chats[chat_id]
            # Se o chat atual foi excluído, mudar para outro chat
            if st.session_state.current_chat_id == chat_id:
                if st.session_state.chats:
                    st.session_state.current_chat_id = next(iter(st.session_state.chats))
                else:
                    # Se não há mais chats, criar um novo
                    new_chat()
            logger.info(f"Chat {chat_id} excluído.")
            save_chats()

    # CSS básico para estilo consistente
    css = """
    <style>
        .centered-title {
            text-align: center;
            padding: 20px 0;
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

    # Inicializar o estado da sessão
    if 'chats' not in st.session_state:
        st.session_state.chats = load_chats()
    if 'current_chat_id' not in st.session_state:
        st.session_state.current_chat_id = next(iter(st.session_state.chats))
    if 'user_interactions' not in st.session_state:
        st.session_state.user_interactions = 0
    if 'total_tokens' not in st.session_state:
        st.session_state.total_tokens = 0
    if 'total_characters' not in st.session_state:
        st.session_state.total_characters = 0
    if 'response_cache' not in st.session_state:
        st.session_state.response_cache = {}
    if 'chat_to_rename' not in st.session_state:
        st.session_state.chat_to_rename = None
    if 'new_chat_title' not in st.session_state:
        st.session_state.new_chat_title = ""
    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = []
    if 'fixed_materials' not in st.session_state:
        # Carregar materiais fixos com indicador de carregamento
        with st.spinner("Carregando materiais..."):
            try:
                materials_text, materials_tokens, materials_chars = load_fixed_materials()
                st.session_state.fixed_materials = materials_text
                logger.info(f"Materiais fixos carregados com sucesso. Total de tokens: {materials_tokens}")
            except Exception as e:
                st.session_state.fixed_materials = f"Erro ao carregar materiais: {e}"
                logger.error(f"Erro ao carregar materiais fixos: {e}")

    # Interface do usuário Streamlit
    col1, col2, col3 = st.columns([1,2,1])

    with col2:
        try:
            logo_path = assets_dir / "LOGO SUCESSO EM VENDAS HORIZONTAL AZUL.png"
            st.image(str(logo_path), width=300, use_container_width=True)
        except FileNotFoundError:
            st.write("Logo não encontrada. Por favor, verifique o caminho da imagem.")

    st.markdown("<h1 class='centered-title'>Consultor I.A. Sucesso em Vendas</h1>", unsafe_allow_html=True)
    st.write("")

    # Barra lateral
    with st.sidebar:
        st.header("Gerenciamento de Chats")
        
        if st.button("Novo Chat", key="btn_new_chat"):
            new_chat()
            st.rerun()
        
        # Upload de arquivos
        st.markdown("### Materiais para este Chat")
        uploaded_files = st.file_uploader(
            "Carregar materiais adicionais (opcional)",
            accept_multiple_files=True,
            type=["pdf", "docx", "json"],
            key="file_uploader"
        )
        
        if uploaded_files:
            st.session_state.uploaded_files = uploaded_files
            # Processar os arquivos carregados
            user_materials, user_tokens, user_chars = process_uploaded_files(uploaded_files)
            # Armazenar os materiais do usuário no chat atual
            st.session_state.chats[st.session_state.current_chat_id]['user_materials'] = user_materials
            st.success(f"Materiais carregados: {len(uploaded_files)} arquivos")
            save_chats()
        
        st.markdown("---")
        st.markdown("### Chats Anteriores")
        
        # Modal para renomear chat
        if st.session_state.chat_to_rename:
            with st.form(key="rename_form"):
                st.text_input("Novo título:", key="new_chat_title", 
                              value=st.session_state.chats[st.session_state.chat_to_rename]['title'])
                col1, col2 = st.columns(2)
                with col1:
                    if st.form_submit_button("Salvar"):
                        rename_chat(st.session_state.chat_to_rename, st.session_state.new_chat_title)
                        st.session_state.chat_to_rename = None
                        st.rerun()
                with col2:
                    if st.form_submit_button("Cancelar"):
                        st.session_state.chat_to_rename = None
                        st.rerun()
        
        # Exibir lista de chats com opções
        for chat_id, chat_data in st.session_state.chats.items():
            col1, col2 = st.columns([5, 1])
            with col1:
                if st.button(f"{chat_data['title']} - {chat_data['date']}", 
                             key=f"chat_btn_{chat_id}",
                             use_container_width=True):
                    st.session_state.current_chat_id = chat_id
                    logger.info(f"Usuário mudou para o chat: {chat_id}")
                    st.rerun()
            with col2:
                # Menu de opções (...)
                if st.button("⋮", key=f"options_{chat_id}"):
                    st.session_state.chat_options_open = chat_id if not hasattr(st.session_state, 'chat_options_open') or st.session_state.chat_options_open != chat_id else None
                    st.rerun()
            
            # Exibir opções se o menu estiver aberto
            if hasattr(st.session_state, 'chat_options_open') and st.session_state.chat_options_open == chat_id:
                option_col1, option_col2 = st.columns(2)
                with option_col1:
                    if st.button("Renomear", key=f"rename_{chat_id}", use_container_width=True):
                        st.session_state.chat_to_rename = chat_id
                        st.session_state.chat_options_open = None
                        st.rerun()
                with option_col2:
                    if st.button("Excluir", key=f"delete_{chat_id}", use_container_width=True):
                        delete_chat(chat_id)
                        st.session_state.chat_options_open = None
                        st.rerun()
        
        # Botão para voltar à página inicial
        st.markdown("---")
        if st.button("← Voltar para a página inicial", key="btn_voltar_consultor"):
            st.session_state.current_app = 'home'
            st.query_params["app"] = "home"
            st.rerun()

    # Adicionando botões de prompt predefinidos
    st.subheader("Prompts Rápidos")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("Vender Produto", key="btn_vender_produto"):
            st.session_state.user_input = ("Me ajude a vender uma (...), preciso de ideias práticas e ações "
                                           "aplicáveis para meu time vender esse produto, preciso que enfatize suas "
                                           "qualidades reais e diferenciais e busque argumentos concisos que "
                                           "naturalmente me ajudem com possíveis objeções.")
    with col2:
        if st.button("Criar Treinamento", key="btn_criar_treinamento"):
            st.session_state.user_input = ("Me ajude a criar um treinamento de (...) com ferramentas e uma lógica de "
                                           "apresentação. Destrinche os tópicos com conteúdos mais práticos e aplicáveis.")
    with col3:
        if st.button("Estratégia de Marketing", key="btn_estrategia_marketing"):
            st.session_state.user_input = ("Preciso de uma estratégia de marketing para aumentar a visibilidade e "
                                           "engajamento do nosso produto. Inclua ideias inovadoras que possam ser "
                                           "implementadas rapidamente e que aproveitem as tendências atuais do mercado.")

    # Inicializar o estado da sessão para a entrada do usuário
    if 'user_input' not in st.session_state:
        st.session_state['user_input'] = ''

    # Modificação na parte do formulário de entrada
    st.write("")
    with st.form(key='input_form', clear_on_submit=True):
        user_input = st.text_input(label='Digite sua mensagem', key='user_input')
        submit_button = st.form_submit_button(label="Enviar")

    # Preparar o contexto completo para este chat
    current_chat = st.session_state.chats[st.session_state.current_chat_id]
    fixed_materials = st.session_state.fixed_materials
    user_materials = current_chat.get('user_materials', '')
    
    # Combinar contexto do agente com materiais fixos e do usuário
    context = f"{agent_context}\n\n"
    
    # Adicionar materiais fixos se existirem
    if fixed_materials:
        context += f"MATERIAIS DE REFERÊNCIA FIXOS:\n{fixed_materials}\n\n"
    
    # Adicionar materiais do usuário se existirem
    if user_materials:
        context += f"MATERIAIS ADICIONADOS PELO USUÁRIO:\n{user_materials}\n\n"
    
    # Adicionar histórico de conversas para contexto
    if current_chat['messages']:
        context += "HISTÓRICO DE CONVERSAS:\n"
        for role, message in current_chat['messages'][-5:]:  # Limitar a 5 mensagens para não sobrecarregar
            context += f"{'Usuário' if role == 'user' else 'Assistente'}: {message}\n"
        context += "\n"

    if submit_button and user_input:
        st.session_state.user_interactions += 1
        logger.info(f"Total de interações do usuário: {st.session_state.user_interactions}")
        
        # Adicionar mensagem do usuário ao histórico
        current_chat['messages'].append(('user', user_input))
        
        # Atualizar o título do chat com base na nova entrada
        if current_chat['title'] == "Novo Chat":
            current_chat['title'] = extract_title(user_input)
        
        # Salvar chats após cada mensagem
        save_chats()
        
        # Gerar resposta
        with st.spinner("Gerando resposta..."):
            response = generate_response(user_input, context)
            
            # Exibir resposta gradualmente
            typing_container = st.empty()
            display_typing_response(response, typing_container)
            
            # Após exibir, remover a resposta da visualização direta
            typing_container.empty()
        
        # Adicionar resposta ao histórico
        current_chat['messages'].append(('agent', response))
        
        # Salvar chats após a resposta
        save_chats()
        
        # Atualizar o contador de tokens e caracteres total
        interaction_tokens = num_tokens_from_string(user_input) + num_tokens_from_string(response)
        interaction_chars = count_characters(user_input) + count_characters(response)
        st.session_state.total_tokens += interaction_tokens
        st.session_state.total_characters += interaction_chars
        logger.info(f"Tokens nesta interação: {interaction_tokens}")
        logger.info(f"Caracteres nesta interação: {interaction_chars}")
        logger.info(f"Total de tokens acumulados: {st.session_state.total_tokens}")
        logger.info(f"Total de caracteres acumulados: {st.session_state.total_characters}")

        # Log de informação sobre o uso do cache
        cache_key = hashlib.md5((user_input + context[:100]).encode()).hexdigest()
        if cache_key in st.session_state.response_cache:
            logger.info("Esta resposta foi recuperada do cache.")
            
        # Recarregar a página para atualizar o histórico
        st.rerun()

        # Exibir histórico do chat atual
    if current_chat['messages']:
        st.subheader("Histórico da Conversa")
        
        # Container para o histórico
        with st.container():
            for role, message in current_chat['messages']:
                if role == 'user':
                    st.info(f"**Você:** {message}")
                else:
                    st.success(f"**Consultor I.A.:** {message}")
    else:
        # Exibir mensagem de boas-vindas quando não há histórico
        st.info("""
        👋 **Bem-vindo ao Consultor I.A. da Sucesso em Vendas!**
        
        Estou aqui para ajudar com suas dúvidas sobre vendas, treinamentos, estratégias de marketing e muito mais.
        
        Use os botões de prompts rápidos acima ou digite sua própria pergunta para começar.
        """)

if __name__ == "__main__":
    app()